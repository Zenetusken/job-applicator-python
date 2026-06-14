#!/usr/bin/env python3
"""Live UI workflow tests for all Tier 1 items.

Tests each Tier 1 change against the real environment:
- vLLM at localhost:8000
- mxbai-embed-large-v1 on GPU
- Real CLI commands
- Real file I/O
"""

from __future__ import annotations

import asyncio
import subprocess
import sys
import tempfile
from pathlib import Path

# Add src to path
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from rich.console import Console
from rich.panel import Panel
from rich.table import Table

console = Console()

# ── Helpers ──────────────────────────────────────────────────────────────────

PASS = "[bold green]PASS[/]"
FAIL = "[bold red]FAIL[/]"
SKIP = "[bold yellow]SKIP[/]"
results: list[tuple[str, str, str]] = []  # (item, test, status)


def report(item: str, test: str, passed: bool, detail: str = ""):
    status = PASS if passed else FAIL
    results.append((item, test, status))
    icon = "✓" if passed else "✗"
    console.print(f"  {icon} {test}" + (f" — {detail}" if detail else ""))


def skip(item: str, test: str, reason: str):
    results.append((item, test, SKIP))
    console.print(f"  ⊘ {test} — SKIP: {reason}")


# ── TIER 1 ITEM A1: Style Analyzer → instructor ─────────────────────────────


async def test_a1_style_analyzer_instructor():
    """Test that style_analyzer uses instructor for structured output."""
    console.print(Panel("[bold]A1: Style Analyzer → instructor response_model[/]", style="cyan"))

    from job_applicator.config import LLMConfig
    from job_applicator.documents.style_analyzer import StyleAnalyzer

    config = LLMConfig()
    analyzer = StyleAnalyzer(config)

    # Create a sample job text for analysis
    job_text = """
    Senior Python Developer at TechCorp

    We are looking for a passionate Senior Python Developer to join our growing team.

    Requirements:
    - 5+ years of Python experience
    - Strong experience with FastAPI or Django
    - PostgreSQL and Redis knowledge
    - Docker and Kubernetes experience

    We offer competitive salary, remote work, and great benefits.
    Join us and help build the future of fintech!
    """

    try:
        guide = await analyzer.analyze(job_text)
        report("A1", "instructor returns StyleGuide", hasattr(guide, "tone"))
        report(
            "A1",
            "guide has key_phrases",
            hasattr(guide, "key_phrases") and len(guide.key_phrases) > 0,
        )
        report(
            "A1",
            "guide has power_words",
            hasattr(guide, "power_words") and len(guide.power_words) > 0,
        )
        report("A1", "guide has tone", isinstance(guide.tone, str) and len(guide.tone) > 0)
        report(
            "A1",
            "guide has sentence_structure",
            hasattr(guide, "sentence_structure") and len(guide.sentence_structure) > 0,
        )
        console.print(f"    tone={guide.tone}")
        console.print(f"    key_phrases={guide.key_phrases[:3]}")
        console.print(f"    power_words={guide.power_words[:3]}")
    except Exception as e:
        report("A1", "instructor analysis", False, str(e)[:200])


# ── TIER 1 ITEM A7: Prompt Version Field ─────────────────────────────────────


def test_a7_prompt_version():
    """Test that prompt_version field exists and defaults to '1.0'."""
    console.print(Panel("[bold]A7: Prompt Version Field[/]", style="cyan"))

    from job_applicator.models import CoverLetterResult, TailoredResume

    # TailoredResume
    tr = TailoredResume(
        original_path="/tmp/resume.txt",
        tailored_text="tailored test",
        job_title="Dev",
        job_company="Corp",
        match_score=0.8,
        semantic_score=0.6,
        skill_score=0.4,
        changes_summary="updated",
    )
    report("A7", "TailoredResume.prompt_version default", tr.prompt_version == "1.0")

    # Try with custom version
    tr2 = TailoredResume(
        original_path="/tmp/resume.txt",
        tailored_text="tailored test",
        job_title="Dev",
        job_company="Corp",
        match_score=0.8,
        semantic_score=0.6,
        skill_score=0.4,
        changes_summary="updated",
        prompt_version="2.0",
    )
    report("A7", "TailoredResume.prompt_version=2.0", tr2.prompt_version == "2.0")

    # CoverLetterResult
    clr = CoverLetterResult(
        job_title="Dev",
        job_company="Corp",
        cover_letter_text="Dear hiring manager...",
    )
    report("A7", "CoverLetterResult.prompt_version default", clr.prompt_version == "1.0")


# ── TIER 1 ITEM B1: Embedding Cache Key ──────────────────────────────────────


def test_b1_cache_key():
    """Test that cache key includes model name and normalize flag, is 32-char MD5."""
    console.print(Panel("[bold]B1: Embedding Cache Key[/]", style="cyan"))

    from job_applicator.config import EmbeddingConfig
    from job_applicator.embeddings.service import EmbeddingService

    config1 = EmbeddingConfig(model_name="model-a", normalize_embeddings=True)
    config2 = EmbeddingConfig(model_name="model-b", normalize_embeddings=True)
    config3 = EmbeddingConfig(model_name="model-a", normalize_embeddings=False)

    svc1 = EmbeddingService(config1)
    svc2 = EmbeddingService(config2)
    svc3 = EmbeddingService(config3)

    key1 = svc1._get_cache_key("hello world")
    key2 = svc2._get_cache_key("hello world")
    key3 = svc3._get_cache_key("hello world")

    report("B1", "key is 32-char MD5 hex", len(key1) == 32, f"len={len(key1)}")
    report("B1", "different model → different key", key1 != key2)
    report("B1", "different normalize → different key", key1 != key3)
    report("B1", "same config → same key", svc1._get_cache_key("hello world") == key1)
    console.print(f"    key1={key1[:16]}...")
    console.print(f"    key2={key2[:16]}...")
    console.print(f"    key3={key3[:16]}...")


# ── TIER 1 ITEM B2: mxbai Query Prefix ──────────────────────────────────────


async def test_b2_query_prefix():
    """Test that mxbai query prefix is applied to resume embeddings."""
    console.print(Panel("[bold]B2: mxbai Query Prefix[/]", style="cyan"))

    from job_applicator.config import EmbeddingConfig
    from job_applicator.embeddings.matching import JobMatcher
    from job_applicator.models import ResumeData

    config = EmbeddingConfig()
    matcher = JobMatcher(config)

    resume = ResumeData(
        raw_text="John Doe\nPython developer\nSkills: Python, FastAPI",
        name="John Doe",
        skills=["Python", "FastAPI"],
        summary="Python developer",
    )

    try:
        emb = matcher.compute_resume_embedding(resume)
        report("B2", "resume embedding computed", emb is not None)
        report("B2", "embedding has correct dimension", len(emb) == 1024, f"dim={len(emb)}")

        # Verify embed_text with prefix works
        emb_prefix = matcher.embed_text(
            "test text", prefix="Represent this sentence for searching relevant passages: "
        )
        emb_no_prefix = matcher.embed_text("test text")
        report("B2", "prefix changes embedding", not all(emb_prefix == emb_no_prefix))
        console.print(f"    embedding sample: [{emb[0]:.4f}, {emb[1]:.4f}, {emb[2]:.4f}, ...]")
    except Exception as e:
        report("B2", "embedding computation", False, str(e)[:100])


# ── TIER 1 ITEM D3: pdftotext -layout flag ──────────────────────────────────


def test_d3_pdftotext_layout():
    """Check if pdftotext -layout flag is being used."""
    console.print(Panel("[bold]D3: pdftotext -layout Flag[/]", style="cyan"))

    # Read the resume.py source to check
    resume_py = Path(__file__).parent.parent / "src" / "job_applicator" / "documents" / "resume.py"
    source = resume_py.read_text()

    uses_layout = "-layout" in source
    report(
        "D3",
        "pdftotext uses -layout flag",
        uses_layout,
        "FOUND in source" if uses_layout else "NOT FOUND — still missing!",
    )


# ── TIER 1 ITEM D7: Semantic/Skill Score Population ──────────────────────────


def test_d7_score_fields():
    """Test that semantic_score and skill_score fields exist and can be populated."""
    console.print(Panel("[bold]D7: Semantic/Skill Score Population[/]", style="cyan"))

    from job_applicator.models import TailoredResume

    tr = TailoredResume(
        original_path="/tmp/resume.txt",
        tailored_text="tailored",
        job_title="Dev",
        job_company="Corp",
        match_score=0.85,
        semantic_score=0.6,
        skill_score=0.4,
        changes_summary="updated",
    )
    report("D7", "semantic_score field works", tr.semantic_score == 0.6)
    report("D7", "skill_score field works", tr.skill_score == 0.4)
    report("D7", "match_score field works", tr.match_score == 0.85)

    # Check that scores can be zero
    tr_default = TailoredResume(
        original_path="/tmp/resume.txt",
        tailored_text="tailored",
        job_title="Dev",
        job_company="Corp",
        match_score=0.5,
        semantic_score=0.0,
        skill_score=0.0,
        changes_summary="updated",
    )
    report("D7", "semantic_score accepts 0.0", tr_default.semantic_score == 0.0)
    report("D7", "skill_score accepts 0.0", tr_default.skill_score == 0.0)


# ── TIER 1 ITEM D1: Python-docx Support ─────────────────────────────────────


def test_d1_docx_support():
    """Test loading a .docx resume file."""
    console.print(Panel("[bold]D1: Python-docx Support[/]", style="cyan"))

    try:
        from docx import Document
    except ImportError:
        skip("D1", "python-docx import", "python-docx not installed")
        return

    from job_applicator.documents.resume import ResumeLoader

    # Create a real .docx file
    doc = Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("jane@example.com")
    doc.add_paragraph("555-9876")
    doc.add_paragraph("")
    doc.add_paragraph("Summary:")
    doc.add_paragraph("Experienced data scientist with ML expertise.")
    doc.add_paragraph("")
    doc.add_paragraph("Skills:")
    doc.add_paragraph("Python, TensorFlow, PyTorch, scikit-learn, Pandas")
    doc.add_paragraph("")
    doc.add_paragraph("Experience:")
    doc.add_paragraph("Data Scientist | AI Corp | 2020-2024")
    doc.add_paragraph("Built ML pipelines processing 1M+ records daily")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        docx_path = Path(f.name)
    doc.save(str(docx_path))

    try:
        loader = ResumeLoader()
        resume = loader.load(docx_path)

        report("D1", "docx loads without error", True)
        report("D1", "name extracted", resume.name == "Jane Smith", f"name={resume.name}")
        report("D1", "email extracted", resume.email == "jane@example.com", f"email={resume.email}")
        report("D1", "skills extracted", len(resume.skills) > 0, f"skills={resume.skills[:3]}")
        report(
            "D1",
            "summary extracted",
            "data scientist" in resume.summary.lower(),
            f"summary={resume.summary[:50]}",
        )
    except Exception as e:
        report("D1", "docx loading", False, str(e)[:100])
    finally:
        docx_path.unlink(missing_ok=True)


# ── TIER 1 ITEM E1: --json Flag ─────────────────────────────────────────────


def test_e1_json_flag():
    """Test that --json flag is available on CLI commands."""
    console.print(Panel("[bold]E1: --json Flag[/]", style="cyan"))

    # Use the installed entry point directly
    result = subprocess.run(
        ["job-applicator", "match", "--help"],
        capture_output=True,
        text=True,
        timeout=10,
        env={
            **subprocess.os.environ,
            "PATH": str(Path(__file__).parent.parent / ".venv" / "bin")
            + ":"
            + subprocess.os.environ.get("PATH", ""),
        },
    )
    match_help = result.stdout

    result2 = subprocess.run(
        ["job-applicator", "search", "--help"],
        capture_output=True,
        text=True,
        timeout=10,
        env={
            **subprocess.os.environ,
            "PATH": str(Path(__file__).parent.parent / ".venv" / "bin")
            + ":"
            + subprocess.os.environ.get("PATH", ""),
        },
    )
    search_help = result2.stdout

    result3 = subprocess.run(
        ["job-applicator", "apply", "--help"],
        capture_output=True,
        text=True,
        timeout=10,
        env={
            **subprocess.os.environ,
            "PATH": str(Path(__file__).parent.parent / ".venv" / "bin")
            + ":"
            + subprocess.os.environ.get("PATH", ""),
        },
    )
    apply_help = result3.stdout

    report("E1", "match has --json flag", "--json" in match_help)
    report("E1", "search has --json flag", "--json" in search_help)
    report("E1", "apply has --json flag", "--json" in apply_help)


# ── TIER 1 ITEM F4: Seniority Detection ─────────────────────────────────────


def test_f4_seniority_detection():
    """Test seniority detection with various job titles."""
    console.print(Panel("[bold]F4: Seniority Detection[/]", style="cyan"))

    from job_applicator.models import JobBoard, JobListing, detect_seniority

    cases = [
        ("Senior Python Developer", "senior"),
        ("Junior Frontend Engineer", "junior"),
        ("Lead Backend Developer", "lead"),
        ("Staff Software Engineer", "staff"),
        ("Principal Architect", "principal"),
        ("Director of Engineering", "director"),
        ("Intern - Summer 2025", "intern"),
        ("Mid-Level DevOps Engineer", "mid"),
        ("Software Engineer", None),
        ("Python Developer", None),
    ]

    for title, expected in cases:
        result = detect_seniority(title)
        passed = result == expected
        report("F4", f"detect('{title}')", passed, f"got={result}, expected={expected}")

    # Test that JobListing accepts seniority field
    job = JobListing(
        title="Senior Python Developer",
        company="Corp",
        url="https://example.com/1",
        board=JobBoard.LINKEDIN,
        seniority="senior",
    )
    report("F4", "JobListing.seniority field", job.seniority == "senior")

    # Test auto-detection via detect_seniority
    job2 = JobListing(
        title="Junior Data Analyst",
        company="Corp",
        url="https://example.com/2",
        board=JobBoard.LINKEDIN,
    )
    detected = detect_seniority(job2.title)
    report("F4", "auto-detect on JobListing", detected == "junior", f"detected={detected}")


# ── TIER 1 ITEM F1/F2: Dependency Cleanup ────────────────────────────────────


def test_f1_f2_dependencies():
    """Test that unused deps are removed and missing deps are added."""
    console.print(Panel("[bold]F1/F2: Dependency Cleanup[/]", style="cyan"))

    # Check pyproject.toml for removed deps (they may still be importable as transitive deps)
    pyproject = Path(__file__).parent.parent / "pyproject.toml"
    content = pyproject.read_text()

    # Extract the [project] dependencies section
    in_deps = False
    deps_section = ""
    for line in content.split("\n"):
        if line.strip().startswith("dependencies"):
            in_deps = True
            continue
        if in_deps:
            if line.strip().startswith("[") or (
                line.strip()
                and not line.strip().startswith('"')
                and not line.strip().startswith("'")
            ):
                if line.strip().startswith("["):
                    break
            deps_section += line + "\n"

    removed = ["httpx", "beautifulsoup4", "pydantic-ai", "crawl4ai"]
    for dep in removed:
        in_deps = dep in deps_section
        report(
            "F1",
            f"{dep} removed from deps",
            not in_deps,
            "still in deps" if in_deps else "confirmed removed",
        )

    # Check that core deps are importable
    core_deps = ["numpy", "docx", "litellm", "instructor", "typer", "rich", "pydantic"]
    for mod in core_deps:
        try:
            __import__(mod)
            report("F2", f"core dep {mod} importable", True)
        except ImportError:
            report("F2", f"core dep {mod} importable", False, "MISSING")

    # Check pyproject.toml has the right deps
    report("F2", "numpy in pyproject.toml", "numpy" in content)
    report("F2", "python-docx in pyproject.toml", "python-docx" in content)


# ── SUMMARY ──────────────────────────────────────────────────────────────────


def print_summary():
    """Print a summary table of all test results."""
    console.print()
    table = Table(title="Tier 1 Live Test Results", show_lines=True)
    table.add_column("Item", style="cyan")
    table.add_column("Test", style="white")
    table.add_column("Status", justify="center")

    passed = sum(1 for _, _, s in results if s == PASS)
    failed = sum(1 for _, _, s in results if s == FAIL)
    skipped = sum(1 for _, _, s in results if s == SKIP)

    for item, test, status in results:
        table.add_row(item, test, status)

    console.print(table)
    console.print(f"\n[bold]Total: {passed} passed, {failed} failed, {skipped} skipped[/]")


# ── MAIN ─────────────────────────────────────────────────────────────────────


async def main():
    console.print(Panel("[bold white]TIER 1 LIVE UI WORKFLOW TESTS[/]", style="blue", expand=False))
    console.print("Environment: vLLM at localhost:8000, GPU available, Python 3.12\n")

    # Non-LLM tests (fast, no GPU needed)
    test_a7_prompt_version()
    test_b1_cache_key()
    test_d3_pdftotext_layout()
    test_d7_score_fields()
    test_d1_docx_support()
    test_e1_json_flag()
    test_f4_seniority_detection()
    test_f1_f2_dependencies()

    # LLM/GPU tests (require live services)
    await test_a1_style_analyzer_instructor()
    await test_b2_query_prefix()

    print_summary()


if __name__ == "__main__":
    asyncio.run(main())
