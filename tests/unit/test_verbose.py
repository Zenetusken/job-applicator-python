"""Tests for global --verbose and --log-file CLI flags."""

from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock

import pytest
from typer.testing import CliRunner

from job_applicator.cli import app

runner = CliRunner()


def test_verbose_flag_is_accepted() -> None:
    result = runner.invoke(app, ["--verbose", "ats-check", "--help"])
    assert result.exit_code == 0


def test_log_file_requires_verbose() -> None:
    result = runner.invoke(app, ["--log-file", "out.json", "ats-check", "--help"])
    assert result.exit_code != 0
    assert "verbose" in result.output.lower()


def test_search_verbose_flag() -> None:
    result = runner.invoke(app, ["--verbose", "search", "--help"])
    assert result.exit_code == 0


def test_config_init_verbose() -> None:
    with runner.isolated_filesystem():
        result = runner.invoke(app, ["--verbose", "config-init", "--output", "config.toml"])
        assert result.exit_code == 0
        assert "Verbose Report" in result.output


def test_json_verbose_keeps_stdout_pure_json(tmp_path: Path) -> None:
    """`--json --verbose`: the verbose observability report renders to STDERR (err_console),
    so stdout stays parseable JSON. Asserts `result.stdout` (stdout-only in click 8.4;
    `result.output` MERGES both streams and so can't see the split) — the report previously
    polluted stdout and broke `<cmd> --json | jq`.
    """
    from docx import Document

    resume = tmp_path / "r.docx"
    doc = Document()
    doc.add_paragraph("Jane Dev\njane@example.com\n(555) 111-2222\nExperience\nSkills: Python")
    doc.save(str(resume))

    result = runner.invoke(app, ["ats-check", "--resume", str(resume), "--json", "--verbose"])
    assert result.exit_code == 0, result.output
    parsed = json.loads(result.stdout)  # raises if the verbose report leaked onto stdout
    assert "score" in parsed
    assert "Verbose Report" in result.stderr  # the report lives on stderr, where logs belong


def test_ats_check_strict_gates_exit_on_incompatible(tmp_path: Path) -> None:
    """`--strict` makes ats-check exit non-zero when the résumé is NOT ATS-compatible (CI
    gating); without it, an incompatible résumé still exits 0 (report-only). The JSON is still
    emitted under `--strict --json` before the non-zero exit."""
    from docx import Document

    resume = tmp_path / "low.docx"
    doc = Document()
    doc.add_paragraph("Jane. I do python things. Hire me.")  # minimal → Not Compatible (~14%)
    doc.save(str(resume))

    r_default = runner.invoke(app, ["ats-check", "--resume", str(resume)])
    assert r_default.exit_code == 0, r_default.output  # report-only: incompatible still exits 0
    assert "Not Compatible" in r_default.output

    r_strict = runner.invoke(app, ["ats-check", "--resume", str(resume), "--strict"])
    assert r_strict.exit_code == 1  # --strict gates the exit on the verdict

    r_json = runner.invoke(app, ["ats-check", "--resume", str(resume), "--strict", "--json"])
    assert r_json.exit_code == 1
    assert json.loads(r_json.stdout)["is_compatible"] is False  # JSON still emitted, then exit 1

    # A COMPATIBLE résumé under --strict must NOT be gated (exit 0) — guard against false-gating.
    good = tmp_path / "good.docx"
    gdoc = Document()
    for line in [
        "Jane Developer",
        "jane@example.com | (555) 123-4567 | San Francisco, CA",
        "Summary",
        "Senior Python engineer, 8 years building async data pipelines and ML services.",
        "Experience",
        "Staff Engineer, Acme Data (2021-Present)",
        "Built async ingestion handling 2B events/day; led a Pydantic v2 + mypy migration.",
        "Education",
        "B.S. Computer Science, State University (2017)",
        "Skills",
        "Python, asyncio, FastAPI, Pydantic, PostgreSQL, Redis, Docker, AWS",
    ]:
        gdoc.add_paragraph(line)
    gdoc.save(str(good))
    r_good = runner.invoke(app, ["ats-check", "--resume", str(good), "--strict"])
    assert r_good.exit_code == 0, r_good.output  # compatible → --strict does not gate


def test_gcl_json_stdout_is_pure_json(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    """`generate-cover-letter --json` emits ONLY the JSON object to stdout — all progress/info
    (Detected tone, Style loaded, the spinner) goes to stderr. Regression guard: 'Detected tone:'
    was printed to stdout, so `gcl --json | jq` broke on the leading non-JSON line. The LIVE
    qa-sanity check needs vLLM; this pins the stdout split in the green unit gate."""
    import job_applicator.cli as cli
    from job_applicator.models import ResumeData

    loader = MagicMock()
    loader.load.return_value = ResumeData(raw_text="Jane\njane@example.com\nPython, SQL")
    gen = MagicMock()
    gen.generate = AsyncMock(return_value="Dear Hiring Manager,\n\nI build async Python systems.")
    tone = MagicMock(primary="professional", confidence=0.9)

    monkeypatch.setattr("job_applicator.documents.resume.ResumeLoader", lambda: loader)
    monkeypatch.setattr(
        "job_applicator.documents.cover_letter.CoverLetterGenerator", lambda *a, **k: gen
    )
    monkeypatch.setattr(cli, "_load_user_profile", lambda settings, *, resume_name="": MagicMock())
    monkeypatch.setattr(cli, "_detect_tone", lambda job: tone)
    monkeypatch.setattr(cli, "_make_runtime", lambda settings: MagicMock())
    monkeypatch.setattr(
        "job_applicator.documents.tone_detector.ToneDetector",
        lambda: MagicMock(format_for_prompt=lambda _tp: ""),
    )

    resume = tmp_path / "r.txt"
    resume.write_text("Jane\njane@example.com\nPython, SQL")
    result = runner.invoke(
        app,
        [
            "generate-cover-letter",
            "-t",
            "Dev",
            "-c",
            "Acme",
            "-d",
            "Python",
            "--resume",
            str(resume),
            "--json",
        ],
    )
    assert result.exit_code == 0, result.output
    parsed = json.loads(result.stdout)  # raises if any progress line leaked onto stdout
    assert parsed["cover_letter"] == "Dear Hiring Manager,\n\nI build async Python systems."
    assert parsed["job_title"] == "Dev"
    assert parsed["company"] == "Acme"
    assert parsed["output_path"]  # default --format txt writes a text artifact
    assert "Detected tone" not in result.stdout  # progress went to stderr
