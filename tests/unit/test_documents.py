"""Unit tests for documents layer."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from pydantic import ValidationError

from job_applicator.config import LLMConfig
from job_applicator.documents.cover_letter import (
    CoverLetterDraft,
    CoverLetterGenerator,
    SourceBackedSentence,
    _application_frame,
    _GeneratedCover,
)
from job_applicator.documents.resume import ResumeLoader
from job_applicator.embeddings.matching import SourceFactRankingResult
from job_applicator.exceptions import DocumentError, LLMError, ResumeNotFoundError
from job_applicator.models import (
    CoverLetterOverlay,
    GroundingReport,
    JobBoard,
    JobListing,
    RankedSourceFact,
    ResumeData,
    SourceFact,
    SourceFactCatalog,
    SourceFactRanking,
    StyleGuide,
    TargetCriteria,
    TargetCriterion,
    UserProfile,
)


def _cover_draft() -> CoverLetterDraft:
    from job_applicator.documents.source_realization import realize_cover_statements

    return CoverLetterDraft(
        body_facts=realize_cover_statements(_cover_source_facts().facts, language="English")
    )


def _source_facts() -> SourceFactCatalog:
    kinds = ["experience", "projects", "education", "experience"]
    return SourceFactCatalog(
        facts=[
            SourceFact(fact_id=f"SRC-{index:03d}", kind=kind, text=f"Source fact {index}.")
            for index, kind in enumerate(kinds, start=1)
        ]
    )


def _cover_source_facts() -> SourceFactCatalog:
    return SourceFactCatalog(facts=_source_facts().facts[:3])


def _cover_ranking_result() -> SourceFactRankingResult:
    facts = _cover_source_facts()
    ranking = SourceFactRanking(
        target_criteria=TargetCriteria(
            job_source_sha256="a" * 64,
            criteria=[TargetCriterion(name="Source evidence", evidence="Source fact")],
        ),
        ranked_facts=[
            RankedSourceFact(
                fact_id=fact.fact_id,
                score=score,
                strongest_similarity=score,
                strongest_criterion_index=0,
            )
            for fact, score in zip(facts.facts, (0.9, 0.8, 0.7), strict=True)
        ],
    )
    return SourceFactRankingResult(facts=facts, ranking=ranking)


def _generated_cover(text: str, resume: ResumeData) -> _GeneratedCover:
    return _GeneratedCover(
        text=text,
        draft=_cover_draft(),
        source_facts=_cover_source_facts(),
        language="English",
    )


def _cover_overlay() -> CoverLetterOverlay:
    return CoverLetterOverlay(
        body_sentences=_cover_draft().body_facts,
        source_body_sha256="a" * 64,
        source_language="en",
    )


def _structured_resume_text(label: str = "Resume text") -> str:
    return f"ALEX MORGAN\n\nSUMMARY\n{label}\n\nEXPERIENCE\nSource experience."


def test_resume_loader_missing_file() -> None:
    loader = ResumeLoader()
    with pytest.raises(ResumeNotFoundError):
        loader.load("/nonexistent/resume.pdf")


def test_resume_loader_unsupported_format(tmp_path: object) -> None:
    import pathlib

    p = pathlib.Path(str(tmp_path)) / "resume.xyz"
    p.write_text("test")
    loader = ResumeLoader()
    with pytest.raises(DocumentError, match="Unsupported"):
        loader.load(p)


def test_resume_loader_text_file(tmp_path: object) -> None:
    import pathlib

    p = pathlib.Path(str(tmp_path)) / "resume.txt"
    p.write_text("John Doe\njohn@example.com\n555-0123\nSkills: Python, FastAPI")
    loader = ResumeLoader()
    resume = loader.load(p)
    assert resume.name == "John Doe"
    assert resume.email == "john@example.com"
    assert "Python" in resume.skills


def test_resume_loader_strips_markdown_bold_from_name(tmp_path: object) -> None:
    """A tailored résumé re-parsed as input should not keep ** around the name."""
    import pathlib

    p = pathlib.Path(str(tmp_path)) / "resume.txt"
    p.write_text("**Alex Rivera**\nalex@example.com\nSkills: Python\n")
    loader = ResumeLoader()
    resume = loader.load(p)
    assert resume.name == "Alex Rivera"


def test_resume_loader_recognizes_technical_skills_header(tmp_path: object) -> None:
    """Parser must recognize qualified skills headers (not just 'Skills')."""
    import pathlib

    p = pathlib.Path(str(tmp_path)) / "resume.txt"
    p.write_text(
        "Jane Smith\njane@example.com\n\n"
        "Technical Skills\nPython\nKubernetes\nTerraform\n\n"
        "Experience\nSenior Engineer at Acme"
    )
    loader = ResumeLoader()
    resume = loader.load(p)
    assert "Python" in resume.skills
    assert "Kubernetes" in resume.skills
    assert "Terraform" in resume.skills


def test_resume_loader_skips_markdown_underline_in_skills(tmp_path: object) -> None:
    """Setext/markdown underline lines under a header must not become a skill."""
    import pathlib

    p = pathlib.Path(str(tmp_path)) / "resume.txt"
    p.write_text(
        "Jane Smith\njane@example.com\n\n"
        "Skills\n------\nPython, FastAPI, PostgreSQL\n\n"
        "Experience\nSenior Engineer at Acme"
    )
    loader = ResumeLoader()
    resume = loader.load(p)
    assert "Python" in resume.skills
    assert "FastAPI" in resume.skills
    assert "PostgreSQL" in resume.skills
    assert "-----" not in resume.skills


def test_resume_loader_skips_markdown_underline_in_summary(tmp_path: object) -> None:
    """Setext/markdown underline lines under a Summary header must not pollute summary."""
    import pathlib

    p = pathlib.Path(str(tmp_path)) / "resume.txt"
    p.write_text(
        "Jane Smith\njane@example.com\n\n"
        "Summary\n-------\nSenior backend engineer with cloud experience.\n\n"
        "Experience\nSenior Engineer at Acme"
    )
    loader = ResumeLoader()
    resume = loader.load(p)
    assert resume.summary.startswith("Senior backend engineer")
    assert "-----" not in resume.summary


def test_summary_fallback_keeps_substantial_first_paragraph() -> None:
    """L-7: with no Summary/Objective header, a substantial first paragraph is the summary."""
    text = (
        "John Doe\njohn@example.com\n555-123-4567\n"
        "Seasoned backend engineer with a decade of experience building "
        "reliable distributed systems for fintech companies.\n\n"
        "Experience\nSenior Engineer at Acme"
    )
    resume = ResumeLoader().parse_text(text)
    assert resume.summary.startswith("Seasoned backend engineer")


def test_summary_fallback_drops_too_short_paragraph() -> None:
    """L-7: a tiny first paragraph is not treated as a summary."""
    text = "John Doe\njohn@example.com\n555-123-4567\nHello there.\n\nExperience\nDev"
    resume = ResumeLoader().parse_text(text)
    assert resume.summary == ""


def test_resume_loader_recognizes_core_competencies_header(tmp_path: object) -> None:
    """'Core Competencies' is a recognized inline skills header."""
    import pathlib

    p = pathlib.Path(str(tmp_path)) / "resume.txt"
    p.write_text(
        "Jane Smith\njane@example.com\n\n"
        "Core Competencies: Python, Docker, AWS\n\n"
        "Experience\nSenior Engineer at Acme"
    )
    loader = ResumeLoader()
    resume = loader.load(p)
    assert "Python" in resume.skills
    assert "Docker" in resume.skills
    assert "AWS" in resume.skills


def test_resume_loader_docx(tmp_path: object) -> None:
    """Test loading a DOCX resume."""
    import pathlib

    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("jane@example.com")
    doc.add_paragraph("Skills: Python, Docker, AWS")
    doc.add_paragraph("Experience: Senior Dev at TechCo")
    p = pathlib.Path(str(tmp_path)) / "resume.docx"
    doc.save(str(p))

    loader = ResumeLoader()
    resume = loader.load(p)
    assert resume.name == "Jane Smith"
    assert resume.email == "jane@example.com"
    assert "Python" in resume.skills


def test_resume_loader_docx_missing_dependency(tmp_path: object) -> None:
    """Test DOCX loading when python-docx is not installed."""
    import pathlib

    p = pathlib.Path(str(tmp_path)) / "resume.docx"
    p.write_bytes(b"fake docx content")

    loader = ResumeLoader()
    with patch.dict("sys.modules", {"docx": None}):
        with pytest.raises(DocumentError, match="python-docx not installed"):
            loader.load(p)


def test_ocr_fallback_triggers_on_short_text(tmp_path: Path) -> None:
    pdf_path = tmp_path / "scanned.pdf"
    pdf_path.write_bytes(b"fake pdf bytes")

    loader = ResumeLoader()
    with (
        patch.object(loader, "_run_pdftotext", return_value=" "),
        patch.object(loader, "_run_pymupdf", return_value="  "),
        patch.object(loader, "_ocr_service", MagicMock()) as mock_ocr,
    ):
        # OCR fallback is triggered, but if both extractors AND OCR yield less
        # than OCR_THRESHOLD chars the loader now raises rather than silently
        # returning an unusable ResumeData.
        mock_ocr.extract_text_from_pdf.return_value = "John Doe\nSkills: Python"
        with pytest.raises(DocumentError, match="insufficient extractable text"):
            loader._load_pdf(pdf_path, ocr_mode="auto")

    mock_ocr.extract_text_from_pdf.assert_called_once_with(pdf_path)


def test_force_ocr_skips_text_extraction(tmp_path: Path) -> None:
    pdf_path = tmp_path / "scanned.pdf"
    pdf_path.write_bytes(b"fake pdf bytes")

    loader = ResumeLoader()
    with (
        patch.object(loader, "_run_pdftotext") as mock_pdftotext,
        patch.object(loader, "_run_pymupdf") as mock_pymupdf,
        patch.object(loader, "_ocr_service", MagicMock()) as mock_ocr,
    ):
        mock_ocr.extract_text_from_pdf.return_value = "OCR text"
        result = loader._load_pdf(pdf_path, ocr_mode="on")

    mock_pdftotext.assert_not_called()
    mock_pymupdf.assert_not_called()
    mock_ocr.extract_text_from_pdf.assert_called_once_with(pdf_path)
    assert result.raw_text == "OCR text"


def test_ocr_mode_off_disables_ocr(tmp_path: Path) -> None:
    pdf_path = tmp_path / "scanned.pdf"
    pdf_path.write_bytes(b"fake pdf bytes")

    long_text = "A" * 150
    loader = ResumeLoader()
    with (
        patch.object(loader, "_run_pdftotext", return_value=long_text),
        patch.object(loader, "_run_pymupdf", return_value=""),
        patch.object(loader, "_ocr_service", MagicMock()) as mock_ocr,
    ):
        result = loader._load_pdf(pdf_path, ocr_mode="off")

    mock_ocr.extract_text_from_pdf.assert_not_called()
    assert long_text in result.raw_text


def test_ocr_mode_off_insufficient_text_raises(tmp_path: Path) -> None:
    pdf_path = tmp_path / "scanned.pdf"
    pdf_path.write_bytes(b"fake pdf bytes")

    loader = ResumeLoader()
    with (
        patch.object(loader, "_run_pdftotext", return_value="X"),
        patch.object(loader, "_run_pymupdf", return_value="X"),
        patch.object(loader, "_ocr_service", MagicMock()) as mock_ocr,
    ):
        with pytest.raises(DocumentError, match="insufficient extractable text"):
            loader._load_pdf(pdf_path, ocr_mode="off")

    mock_ocr.extract_text_from_pdf.assert_not_called()


def test_ocr_failure_falls_back_to_extracted_text(tmp_path: Path) -> None:
    pdf_path = tmp_path / "scanned.pdf"
    pdf_path.write_bytes(b"fake pdf bytes")

    long_text = "A" * 150
    loader = ResumeLoader()
    with (
        patch.object(loader, "_run_pdftotext", return_value="short"),
        patch.object(loader, "_run_pymupdf", return_value=long_text),
        patch.object(loader, "_ocr_service", MagicMock()) as mock_ocr,
    ):
        mock_ocr.extract_text_from_pdf.side_effect = DocumentError("OCR failed")
        result = loader._load_pdf(pdf_path, ocr_mode="auto")

    assert long_text in result.raw_text


def test_ocr_failure_with_no_text_raises(tmp_path: Path) -> None:
    pdf_path = tmp_path / "scanned.pdf"
    pdf_path.write_bytes(b"fake pdf bytes")

    loader = ResumeLoader()
    with (
        patch.object(loader, "_run_pdftotext", return_value=""),
        patch.object(loader, "_run_pymupdf", return_value=""),
        patch.object(loader, "_ocr_service", MagicMock()) as mock_ocr,
    ):
        mock_ocr.extract_text_from_pdf.side_effect = DocumentError("OCR failed")
        with pytest.raises(DocumentError):
            loader._load_pdf(pdf_path, ocr_mode="auto")


def test_image_resume_uses_ocr(tmp_path: Path) -> None:
    img_path = tmp_path / "resume.png"
    from PIL import Image

    Image.new("RGB", (50, 50), color="white").save(img_path)

    loader = ResumeLoader()
    with patch.object(loader, "_ocr_service", MagicMock()) as mock_ocr:
        mock_ocr.extract_text_from_image.return_value = "OCR text"
        result = loader.load(img_path, ocr_mode="on")

    mock_ocr.extract_text_from_image.assert_called_once_with(img_path)
    assert result.raw_text == "OCR text"


def test_force_ocr_failure_raises(tmp_path: Path) -> None:
    pdf_path = tmp_path / "scanned.pdf"
    pdf_path.write_bytes(b"fake pdf bytes")

    loader = ResumeLoader()
    with patch.object(loader, "_ocr_service", MagicMock()) as mock_ocr:
        mock_ocr.extract_text_from_pdf.side_effect = DocumentError("OCR failed")
        with pytest.raises(DocumentError):
            loader._load_pdf(pdf_path, ocr_mode="on")


def test_parse_text_records_confidence_and_method() -> None:
    loader = ResumeLoader()
    text = "John Doe\njohn@example.com\n555-0123\nSkills: Python\nExperience\nEducation"
    resume = loader.parse_text(text, method="text")
    assert resume.parse_method == "text"
    assert resume.parse_confidence > 0.0
    assert resume.parse_confidence <= 1.0


def test_compute_confidence_empty_text() -> None:
    loader = ResumeLoader()
    assert loader._compute_confidence("") == 0.0
    assert loader._compute_confidence("   ") == 0.0


def test_compute_confidence_increases_with_signals() -> None:
    loader = ResumeLoader()
    base = loader._compute_confidence("John Doe\n")
    with_contact = loader._compute_confidence("John Doe\njohn@example.com\n555-0123\n")
    with_sections = loader._compute_confidence(
        "John Doe\njohn@example.com\n555-0123\nSkills: Python\nExperience\nEducation\n"
    )
    assert with_contact > base
    assert with_sections > with_contact


def test_pdf_consensus_selects_best_parser(tmp_path: Path) -> None:
    pdf_path = tmp_path / "resume.pdf"
    pdf_path.write_bytes(b"fake pdf bytes")

    loader = ResumeLoader()
    good_text = "A" * 200 + "\njohn@example.com\nSkills: Python\nExperience\nEducation"
    bad_text = "garbled"
    with (
        patch.object(loader, "_run_pdftotext", return_value=bad_text),
        patch.object(loader, "_run_pymupdf", return_value=good_text),
    ):
        result = loader._load_pdf(pdf_path, ocr_mode="off")

    assert result.raw_text == good_text
    assert result.parse_method == "pymupdf"
    assert result.parse_confidence > 0.5


def test_password_protected_pdf_raises(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    pdf_path = tmp_path / "locked.pdf"
    pdf_path.write_bytes(b"fake pdf bytes")

    def locked(*_args: object, **_kwargs: object) -> None:
        raise RuntimeError("Document is password protected")

    monkeypatch.setattr("fitz.open", locked)
    loader = ResumeLoader()
    with pytest.raises(DocumentError, match="password-protected"):
        loader._load_pdf(pdf_path, ocr_mode="off")


def test_resume_loader_wraps_unreadable_docx_as_document_error(tmp_path: Path) -> None:
    """A corrupt/empty .docx surfaces as a typed DocumentError via load() — python-docx's
    PackageNotFoundError must not leak as a traceback (the loader's contract is
    JobApplicatorError-only)."""
    bad = tmp_path / "empty.docx"
    bad.write_bytes(b"")
    with pytest.raises(DocumentError):
        ResumeLoader().load(bad)


def test_resume_loader_corrupt_pdf_is_document_error(tmp_path: Path) -> None:
    """A corrupt .pdf surfaces as a typed DocumentError — here via the PDF consensus
    'insufficient extractable text' path (ocr off). Guards the user-facing contract that
    load() raises only JobApplicatorError; the load() try/except wrapper itself is
    exercised by the corrupt-.docx test above (python-docx PackageNotFoundError → DocumentError).
    """
    bad = tmp_path / "corrupt.pdf"
    bad.write_bytes(b"not a real pdf %%\x00\x01garbage")
    with pytest.raises(DocumentError):
        ResumeLoader().load(bad, ocr_mode="off")


def test_resume_loader_empty_text_is_document_error(tmp_path: Path) -> None:
    """A VALID file with no extractable text (empty/whitespace) → typed DocumentError
    ('no extractable text'), not a misleading 0.14 ATS 'score' on nothing (QA pass-2 B3)."""
    blank = tmp_path / "blank.txt"
    blank.write_text("   \n\t  ", encoding="utf-8")
    with pytest.raises(DocumentError, match="no extractable text"):
        ResumeLoader().load(blank)


def test_resume_loader_directory_path_names_the_target(tmp_path: Path) -> None:
    """A directory (no extension) → DocumentError that NAMES the path, not an empty
    'Unsupported resume format: ' message."""
    with pytest.raises(DocumentError) as ei:
        ResumeLoader().load(tmp_path)
    assert tmp_path.name in str(ei.value)


def _validation_job_and_resume() -> tuple[JobListing, ResumeData]:
    """Return a job/company that does NOT appear on the resume."""
    job = JobListing(
        title="Python Dev", company="Acme", url="https://example.com/1", board=JobBoard.LINKEDIN
    )
    resume = ResumeData(raw_text="John Doe\nBackend engineer at OtherCorp")
    return job, resume


def test_cover_letter_validation_rejects_empty() -> None:
    config = LLMConfig()
    generator = CoverLetterGenerator(config)
    _job, resume = _validation_job_and_resume()
    with pytest.raises(LLMError, match="empty"):
        generator._validate_output(_generated_cover("   ", resume), resume=resume)


def test_cover_letter_validation_rejects_too_short() -> None:
    config = LLMConfig()
    generator = CoverLetterGenerator(config)
    _job, resume = _validation_job_and_resume()
    with pytest.raises(LLMError, match="too short"):
        generator._validate_output(_generated_cover("Sincerely,\nJohn Doe", resume), resume=resume)


def test_cover_letter_validation_rejects_placeholders() -> None:
    config = LLMConfig()
    generator = CoverLetterGenerator(config)
    _job, resume = _validation_job_and_resume()
    with pytest.raises(LLMError, match="placeholder"):
        generator._validate_output(
            _generated_cover(
                "Dear [Hiring Manager],\n\nBody text that is long enough to pass the length check. "
                "It keeps going so the validator does not reject it for being too short.\n\n"
                "Sincerely,\nJohn Doe",
                resume,
            ),
            resume=resume,
        )


def test_cover_letter_appends_sign_off_without_rewriting_body() -> None:
    user = UserProfile(first_name="Jane", last_name="Roe", email="j@e.com", phone="")
    body = "Application opening.\n\nSource-backed body.\n\nDiscussion request."
    out = CoverLetterGenerator._append_sign_off(body, user)
    assert out.startswith(body)
    assert out.rstrip().endswith("Sincerely,\nJane Roe")


def test_cover_letter_appends_french_sign_off() -> None:
    user = UserProfile(first_name="Alex", last_name="Morgan", email="a@e.com", phone="")
    text = "Ouverture.\n\nFaits.\n\nDemande de discussion."
    out = CoverLetterGenerator._append_sign_off(text, user, "French")
    assert out.endswith("Cordialement,\nAlex Morgan")
    assert "Sincerely" not in out


def test_french_application_frame_is_independent_of_title_and_evidence_inflection() -> None:
    job = JobListing(
        title="Analyste réseau principal(e)",
        company="Intact",
        url="https://example.test/job",
        board=JobBoard.INDEED,
    )

    opening, closing = _application_frame(job, "French")

    assert opening.startswith("Je vous présente ma candidature chez Intact")
    assert "poste suivant : Analyste réseau principal(e)" in opening
    assert "directement de mon parcours" in opening
    assert "poste de Analyste" not in opening
    assert "mon projets" not in opening
    assert closing.startswith("Ensemble, ces exemples")


def test_cover_letter_quality_accepts_three_french_body_paragraphs() -> None:
    from job_applicator.documents.quality_eval import assess_cover_letter

    text = (
        "Je souhaite postuler au poste et présenter mon parcours à votre équipe. "
        "Mon parcours combine une formation en cybersécurité opérationnelle avec des "
        "laboratoires pratiques en SIEM, SOC et réponse aux incidents.\n\n"
        "Dans mon cours à Northbridge Technical Institute, j'ai participé à des laboratoires "
        "pratiques sur le SIEM, les opérations SOC, la détection d'intrusion, l'EDR et la "
        "réponse aux incidents. Mon expérience en support technique m'a appris à communiquer "
        "clairement avec les équipes techniques.\n\n"
        "Je suis convaincu que mon profil correspond bien au poste et je serais ravi de "
        "discuter de ma candidature en détail.\n\n"
        "Cordialement,\nALEX MORGAN"
    )

    report = assess_cover_letter(text, applicant_name="ALEX MORGAN")

    assert "cover letter should usually have three focused body paragraphs" not in report.warnings


def test_cover_letter_generator_template() -> None:
    config = LLMConfig()
    generator = CoverLetterGenerator(config)
    job = JobListing(
        title="Python Dev",
        company="Acme",
        url="https://example.com/1",
        board=JobBoard.LINKEDIN,
    )
    user = UserProfile(
        first_name="John",
        last_name="Doe",
        email="john@example.com",
        phone="555-0123",
    )
    resume = ResumeData(raw_text="test", skills=["Python"])
    letter = generator.generate_from_template(job, user, resume)
    assert "Acme" in letter
    assert "Python Dev" in letter
    assert "John" in letter


def test_cover_letter_output_model() -> None:
    from job_applicator.documents.cover_letter import CoverLetterOutput

    output = CoverLetterOutput(
        cover_letter="Dear Hiring Manager, ...",
        key_points=["Python experience", "FastAPI expertise"],
    )
    assert output.cover_letter.startswith("Dear")
    assert len(output.key_points) == 2


def test_cover_letter_generator_has_no_applicant_prose_completion_stage() -> None:
    generator = CoverLetterGenerator(LLMConfig())

    assert not hasattr(generator, "_structured_completion")


def test_ocr_service_extracts_text_from_image(tmp_path: Path) -> None:
    from job_applicator.documents.ocr import OCRService

    service = OCRService()
    # PaddleOCR is lazy-loaded; mock it to avoid heavy model init in unit tests.
    service._ocr = MagicMock()
    service._ocr.ocr.return_value = [[([[0, 0], [10, 0], [10, 10], [0, 10]], ("Hello", 0.99))]]

    img_path = tmp_path / "resume.png"
    # Create a tiny blank PNG using PIL
    from PIL import Image

    Image.new("RGB", (50, 50), color="white").save(img_path)

    text = service.extract_text_from_image(img_path)
    assert "Hello" in text


@pytest.fixture
def llm_config() -> LLMConfig:
    return LLMConfig(api_base="http://localhost:8000/v1", model="test-model")


class TestCoverLetterWithTone:
    @pytest.mark.asyncio
    async def test_generate_excludes_tone_power_words_from_fact_prompt(
        self, llm_config: LLMConfig
    ) -> None:
        generator = CoverLetterGenerator(llm_config)
        generator._cover_evidence_candidates = MagicMock(  # type: ignore[method-assign]
            return_value=_source_facts()
        )
        generator._select_source_facts = AsyncMock(  # type: ignore[method-assign]
            return_value=_cover_ranking_result()
        )

        job = JobListing(
            title="Dev",
            company="Co",
            url="https://example.com",
            board=JobBoard.INDEED,
        )
        user = UserProfile(first_name="John", last_name="Doe", email="j@e.com", phone="123")
        resume = ResumeData(raw_text=_structured_resume_text(), skills=["Python"])

        await generator.generate(
            job,
            user,
            resume,
            tone_section="TONE: Corporate\n- Power words: leveraged",
        )

        call_args = generator._select_source_facts.call_args  # type: ignore[attr-defined]
        assert "TONE: Corporate" not in str(call_args)
        assert "leveraged" not in str(call_args)

    @pytest.mark.asyncio
    async def test_generate_does_not_use_tailored_resume_as_source(
        self, llm_config: LLMConfig
    ) -> None:
        generator = CoverLetterGenerator(llm_config)
        generator._cover_evidence_candidates = MagicMock(  # type: ignore[method-assign]
            return_value=_source_facts()
        )
        generator._select_source_facts = AsyncMock(  # type: ignore[method-assign]
            return_value=_cover_ranking_result()
        )

        job = JobListing(
            title="Dev",
            company="Co",
            url="https://example.com",
            board=JobBoard.INDEED,
        )
        user = UserProfile(first_name="John", last_name="Doe", email="j@e.com", phone="123")
        resume = ResumeData(raw_text=_structured_resume_text("Original resume"), skills=["Python"])
        tailored = "Tailored resume with optimized Python experience"

        await generator.generate(
            job,
            user,
            resume,
            tailored_resume_text=tailored,
        )

        call_args = generator._select_source_facts.call_args  # type: ignore[attr-defined]
        assert "Tailored resume with optimized" not in str(call_args)

    @pytest.mark.asyncio
    async def test_refine_routes_feedback_to_selection_only(self) -> None:
        config = LLMConfig(api_base="http://localhost:8000/v1", model="m", max_tokens=1234)
        generator = CoverLetterGenerator(config)
        generator._cover_evidence_candidates = MagicMock(  # type: ignore[method-assign]
            return_value=_source_facts()
        )
        generator._select_source_facts = AsyncMock(  # type: ignore[method-assign]
            return_value=_cover_ranking_result()
        )

        job = JobListing(
            title="Dev",
            company="Co",
            url="https://example.com",
            board=JobBoard.INDEED,
        )
        user = UserProfile(first_name="John", last_name="Doe", email="j@e.com", phone="123")
        resume = ResumeData(raw_text=_structured_resume_text(), skills=["Python"])

        await generator.refine(
            job,
            user,
            resume,
            current_text="Old cover letter.",
            user_feedback="Emphasize incident response.",
        )

        assert generator._select_source_facts.call_args.kwargs["selection_focus"] == (
            "Emphasize incident response."
        )


def test_password_protected_pdf_detected_via_needs_pass(tmp_path: Path) -> None:
    """M3: PyMuPDF opens an encrypted PDF without raising; needs_pass is the gate."""
    pdf_path = tmp_path / "locked.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 fake")
    loader = ResumeLoader()
    fake_doc = MagicMock()
    fake_doc.needs_pass = True
    fake_fitz = MagicMock()
    fake_fitz.open.return_value = fake_doc
    with patch.dict("sys.modules", {"fitz": fake_fitz}):
        assert loader._is_password_protected(pdf_path) is True
    fake_doc.close.assert_called_once()


def test_unprotected_pdf_not_flagged_as_password_protected(tmp_path: Path) -> None:
    """M3: a normal PDF (needs_pass False) is not flagged as protected."""
    pdf_path = tmp_path / "ok.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 fake")
    loader = ResumeLoader()
    fake_doc = MagicMock()
    fake_doc.needs_pass = False
    fake_fitz = MagicMock()
    fake_fitz.open.return_value = fake_doc
    with patch.dict("sys.modules", {"fitz": fake_fitz}):
        assert loader._is_password_protected(pdf_path) is False


def test_confidence_counts_headers_not_midline_mentions() -> None:
    """Confidence counts line-anchored section HEADERS, not bare substrings."""
    loader = ResumeLoader()
    headers = "Experience\nEducation\nSkills\nProjects\nCertifications"
    prose = "my experience education skills projects certifications summary blah"
    assert loader._compute_confidence(headers) > loader._compute_confidence(prose)


def test_confidence_vocab_aligns_with_skills_extractor() -> None:
    """Confidence credits 'Core Competencies' (aligned with _extract_skills_section)."""
    loader = ResumeLoader()
    with_header = "Jane Doe\n\nCore Competencies\nPython, SQL"
    without = "Jane Doe\n\nrandom prose mentioning competencies in a sentence"
    assert loader._compute_confidence(with_header) > loader._compute_confidence(without)


def test_phone_extraction_rejects_year_runs() -> None:
    """Phone: a run of years must not be mistaken for a phone number."""
    loader = ResumeLoader()
    assert loader._extract_phone("Worked 2019 2020 2021 2022 2023 at Acme") == ""


def test_phone_extraction_accepts_real_phone() -> None:
    """Phone: normally-formatted numbers are still extracted."""
    loader = ResumeLoader()
    assert "555" in loader._extract_phone("Call me at 555-123-4567")
    assert loader._extract_phone("Reach me: +1 (555) 123-4567").strip() != ""


def test_cover_letter_source_citations_reject_unknown_ids() -> None:
    draft = CoverLetterDraft(
        body_facts=[
            SourceBackedSentence(text="Body fact one.", fact_ids=["SRC-999"]),
            SourceBackedSentence(text="Body fact two.", fact_ids=["SRC-001"]),
            SourceBackedSentence(text="Body fact three.", fact_ids=["SRC-002"]),
        ],
    )
    generator = CoverLetterGenerator(LLMConfig())
    source_facts = _cover_source_facts()

    with pytest.raises(LLMError, match="unknown source fact IDs"):
        generator._validate_source_fact_citations(draft, source_facts)


async def test_cover_letter_fact_selection_is_deterministic_and_source_only() -> None:
    matcher = MagicMock()
    matcher.rank_source_facts = AsyncMock(return_value=_cover_ranking_result())
    generator = CoverLetterGenerator(LLMConfig(), matcher=matcher)
    job, _user, _resume = _cl_inputs()

    selected = await generator._select_source_facts(job, _source_facts())

    assert [fact.fact_id for fact in selected.facts.facts] == [
        "SRC-001",
        "SRC-002",
        "SRC-003",
    ]
    assert all(fact in _source_facts().facts for fact in selected.facts.facts)


def test_cover_letter_source_citations_reject_non_deterministic_claim() -> None:
    from job_applicator.documents.source_realization import realize_cover_statements

    generator = CoverLetterGenerator(LLMConfig())
    source_facts = _cover_source_facts()
    statements = realize_cover_statements(source_facts.facts, language="English")
    statements[1] = statements[1].model_copy(
        update={"text": statements[1].text + " Ensured success."}
    )
    draft = CoverLetterDraft(body_facts=statements)

    with pytest.raises(LLMError, match="differs from deterministic source realization"):
        generator._validate_source_fact_citations(draft, source_facts)


def test_cover_letter_realization_excludes_targeting_context() -> None:
    from job_applicator.documents.source_realization import realize_cover_statements

    source_facts = _cover_source_facts()
    source_facts.facts[0] = source_facts.facts[0].model_copy(
        update={"context": "Certificate in Progress | Metro College"}
    )
    generator = CoverLetterGenerator(LLMConfig(language="en"))
    draft = CoverLetterDraft(
        body_facts=realize_cover_statements(source_facts.facts, language="English")
    )
    generator._validate_source_fact_citations(draft, source_facts)

    assert all("Certificate in Progress" not in sentence.text for sentence in draft.body_facts)


def test_cover_letter_source_citations_require_each_selected_fact_once() -> None:
    source_facts = _cover_source_facts()
    draft = CoverLetterDraft(
        body_facts=[
            SourceBackedSentence(text="First fact.", fact_ids=["SRC-001"]),
            SourceBackedSentence(text="First fact again.", fact_ids=["SRC-001"]),
            SourceBackedSentence(text="Second fact.", fact_ids=["SRC-002"]),
        ]
    )
    generator = CoverLetterGenerator(LLMConfig())

    with pytest.raises(LLMError, match="use each selected source fact exactly once"):
        generator._validate_source_fact_citations(draft, source_facts)


def test_source_backed_sentence_requires_evidence() -> None:
    from job_applicator.documents.cover_letter import SourceBackedSentence

    with pytest.raises(ValidationError):
        SourceBackedSentence(text="Unsupported sentence.", fact_ids=[])


def _cl_inputs() -> tuple[JobListing, UserProfile, ResumeData]:
    job = JobListing(
        title="Dev",
        company="Co",
        url="https://e.com",
        requirements=["Python"],
        board=JobBoard.INDEED,
    )
    user = UserProfile(first_name="J", last_name="D", email="j@e.com", phone="1")
    resume = ResumeData(raw_text=_structured_resume_text(), skills=["Python"])
    return job, user, resume


@pytest.mark.asyncio
async def test_generate_verified_delegates_to_inline_validated_generation() -> None:
    gen = CoverLetterGenerator(LLMConfig(model="m"))
    gen.generate_verified_with_overlay = AsyncMock(  # type: ignore[method-assign]
        return_value=("LETTER A", _cover_overlay())
    )

    assert await gen.generate_verified(*_cl_inputs()) == "LETTER A"
    gen.generate_verified_with_overlay.assert_awaited_once()  # type: ignore[attr-defined]


async def test_refine_verified_returns_letter_and_report() -> None:
    gen = CoverLetterGenerator(LLMConfig(model="m"))
    gen.refine_verified_with_overlay = AsyncMock(  # type: ignore[method-assign]
        return_value=("REFINED LETTER", _cover_overlay(), GroundingReport())
    )
    job, user, resume = _cl_inputs()
    letter, report = await gen.refine_verified(job, user, resume, "current", "make it formal")
    assert letter == "REFINED LETTER"
    assert report is not None and report.clean
    gen.refine_verified_with_overlay.assert_awaited_once()  # type: ignore[attr-defined]


async def test_refine_appends_sign_off() -> None:
    """Refined letters get the same deterministic sign-off guarantee as first drafts."""
    gen = CoverLetterGenerator(LLMConfig(model="m"))
    gen._cover_evidence_candidates = MagicMock(  # type: ignore[method-assign]
        return_value=_source_facts()
    )
    gen._select_source_facts = AsyncMock(  # type: ignore[method-assign]
        return_value=_cover_ranking_result()
    )

    job, user, resume = _cl_inputs()
    letter = await gen.refine(job, user, resume, "current letter", "make it tighter")

    assert letter.endswith("Sincerely,\nJ D")


# --- Multi-file style-guide loader --------------------------------------------------


class TestLoadStyleGuide:
    """Cycle 2b: CoverLetterGenerator.load_style_guide is the single shared helper
    used by apply, batch, tailor, and generate-cover-letter."""

    @pytest.fixture
    def config(self) -> LLMConfig:
        return LLMConfig(api_base="http://localhost:8000/v1", model="test-model")

    @pytest.fixture
    def generator(self, config: LLMConfig) -> CoverLetterGenerator:
        return CoverLetterGenerator(config)

    @pytest.fixture
    def style(self) -> StyleGuide:
        return StyleGuide(
            tone="professional",
            sentence_structure="varied",
            vocabulary_level="technical",
            paragraph_style="clear",
            formatting_notes="",
            sample_paragraph="",
        )

    @pytest.mark.asyncio
    async def test_single_text_file_analyzed(
        self,
        generator: CoverLetterGenerator,
        tmp_path: Path,
        style: StyleGuide,
    ) -> None:
        path = tmp_path / "style.txt"
        path.write_text("Professional cover letter tone example.", encoding="utf-8")
        resume = ResumeData(
            raw_text="Professional cover letter tone example.", name="", email="", skills=[]
        )

        with (
            patch(
                "job_applicator.documents.cover_letter.ResumeLoader.load",
                return_value=resume,
            ) as mock_load,
            patch(
                "job_applicator.documents.cover_letter.StyleAnalyzer.analyze",
                new_callable=AsyncMock,
                return_value=style,
            ) as mock_analyze,
        ):
            result = await generator.load_style_guide(str(path))

        assert result is style
        mock_load.assert_called_once_with(path, ocr_mode="auto")
        mock_analyze.assert_awaited_once_with("Professional cover letter tone example.")

    @pytest.mark.asyncio
    async def test_single_pdf_loaded_via_resume_loader(
        self,
        generator: CoverLetterGenerator,
        tmp_path: Path,
        style: StyleGuide,
    ) -> None:
        pdf_path = tmp_path / "style.pdf"
        pdf_path.write_text("fake pdf", encoding="utf-8")
        resume = ResumeData(raw_text="PDF style text", name="", email="", skills=[])

        with (
            patch(
                "job_applicator.documents.cover_letter.ResumeLoader.load",
                return_value=resume,
            ) as mock_load,
            patch(
                "job_applicator.documents.cover_letter.StyleAnalyzer.analyze",
                new_callable=AsyncMock,
                return_value=style,
            ) as mock_analyze,
        ):
            result = await generator.load_style_guide(str(pdf_path), ocr_mode="on")

        assert result is style
        mock_load.assert_called_once_with(pdf_path, ocr_mode="on")
        mock_analyze.assert_awaited_once_with("PDF style text")

    @pytest.mark.asyncio
    async def test_multiple_text_files_use_analyze_multiple(
        self,
        generator: CoverLetterGenerator,
        tmp_path: Path,
        style: StyleGuide,
    ) -> None:
        p1 = tmp_path / "a.txt"
        p2 = tmp_path / "b.txt"
        p1.write_text("Tone A", encoding="utf-8")
        p2.write_text("Tone B", encoding="utf-8")

        def _fake_load(path: Path, ocr_mode: str = "auto") -> ResumeData:
            text = "Tone A" if path.name == "a.txt" else "Tone B"
            return ResumeData(raw_text=text, name="", email="", skills=[])

        with (
            patch(
                "job_applicator.documents.cover_letter.ResumeLoader.load",
                side_effect=_fake_load,
            ) as mock_load,
            patch(
                "job_applicator.documents.cover_letter.StyleAnalyzer.analyze_multiple",
                new_callable=AsyncMock,
                return_value=style,
            ) as mock_multiple,
        ):
            result = await generator.load_style_guide(f"{p1}, {p2}")

        assert result is style
        assert mock_load.call_count == 2
        mock_multiple.assert_awaited_once_with(["Tone A", "Tone B"])

    @pytest.mark.asyncio
    async def test_mixed_pdf_and_text_use_analyze_multiple(
        self,
        generator: CoverLetterGenerator,
        tmp_path: Path,
        style: StyleGuide,
    ) -> None:
        txt = tmp_path / "a.txt"
        pdf = tmp_path / "b.pdf"
        txt.write_text("Text tone", encoding="utf-8")
        pdf.write_text("fake", encoding="utf-8")

        def _fake_load(path: Path, ocr_mode: str = "auto") -> ResumeData:
            text = "Text tone" if path.name == "a.txt" else "PDF tone"
            return ResumeData(raw_text=text, name="", email="", skills=[])

        with (
            patch(
                "job_applicator.documents.cover_letter.ResumeLoader.load",
                side_effect=_fake_load,
            ) as mock_load,
            patch(
                "job_applicator.documents.cover_letter.StyleAnalyzer.analyze_multiple",
                new_callable=AsyncMock,
                return_value=style,
            ) as mock_multiple,
        ):
            result = await generator.load_style_guide(f"{txt},{pdf}")

        assert result is style
        assert mock_load.call_count == 2
        mock_multiple.assert_awaited_once_with(["Text tone", "PDF tone"])

    @pytest.mark.asyncio
    async def test_missing_file_raises(self, generator: CoverLetterGenerator) -> None:
        with pytest.raises(DocumentError, match="not found"):
            await generator.load_style_guide("/nonexistent/style.txt")

    @pytest.mark.asyncio
    async def test_empty_string_raises(self, generator: CoverLetterGenerator) -> None:
        with pytest.raises(DocumentError, match="No style guide paths"):
            await generator.load_style_guide("  ,  ")
