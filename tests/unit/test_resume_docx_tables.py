"""DOCX résumé parsing must extract TABLE cells, not just paragraphs (audit AI-H5).

Real résumés routinely place the contact header and the skills section in tables. The prior
``_load_docx`` read only ``doc.paragraphs``, so that content was silently dropped from
``raw_text`` — the email, phone, and skills never reached matching/tailoring, and a tailored
CV could go out missing the candidate's contact block. These tests pin that table content is
recovered AND placed in document order (the contact table is first, so the name heuristic
finds the name, not the section header that used to lead the text).
"""

from __future__ import annotations

from pathlib import Path

from job_applicator.documents.resume import ResumeLoader


def _make_table_resume(path: Path) -> None:
    """A résumé with the contact header and skills in TABLES (as real ones often are)."""
    from docx import Document

    doc = Document()
    # Contact header in a table — the FIRST block in the document.
    contact = doc.add_table(rows=1, cols=4)
    contact.cell(0, 0).text = "Jane Doe"
    contact.cell(0, 1).text = "jane.doe@example.com"
    contact.cell(0, 2).text = "514-555-0199"
    contact.cell(0, 3).text = "Montréal, QC"
    # Experience as paragraphs (these already worked).
    doc.add_paragraph("Experience")
    doc.add_paragraph("Security Analyst")
    doc.add_paragraph("ACME Corp, Montréal, QC")
    doc.add_paragraph("2020 - Present")
    # Skills in a table (one skill per cell).
    doc.add_paragraph("Skills")
    skills = doc.add_table(rows=2, cols=2)
    skills.cell(0, 0).text = "Python"
    skills.cell(0, 1).text = "SIEM"
    skills.cell(1, 0).text = "Incident Response"
    skills.cell(1, 1).text = "TCP/IP"
    doc.save(str(path))


def test_load_docx_extracts_contact_and_skills_from_tables(tmp_path: Path) -> None:
    p = tmp_path / "resume.docx"
    _make_table_resume(p)

    data = ResumeLoader().load(str(p))

    # Contact (was in a table → previously dropped from raw_text entirely):
    assert data.email == "jane.doe@example.com"
    assert "514-555-0199" in data.raw_text
    assert "Jane Doe" in data.raw_text
    # Skills (in a table → previously dropped):
    for skill in ("Python", "SIEM", "Incident Response", "TCP/IP"):
        assert skill in data.raw_text, f"{skill!r} missing from extracted text"
    # Document ORDER: the contact table is first, so the name heuristic finds the name,
    # not the "Experience" header that used to lead the paragraph-only text.
    assert data.raw_text.index("Jane Doe") < data.raw_text.index("Experience")
    assert data.name == "Jane Doe"


def test_load_docx_skips_malformed_table_best_effort(tmp_path: Path) -> None:
    """A structurally malformed table (an orphan ``vMerge`` continuation — a real export
    artifact from Google Docs / LibreOffice / Word) makes python-docx's grid walk raise. It must
    NOT fail the whole résumé: the paragraphs are still extracted and the bad table is skipped
    with a logged warning (best-effort + disclosed — not a fabricated default). Regression guard:
    before table extraction this file parsed fine (paragraphs never touched the table)."""
    from docx import Document
    from docx.oxml.ns import qn

    p = tmp_path / "malformed.docx"
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane@example.com")
    bad = doc.add_table(rows=2, cols=2)
    bad.cell(0, 0).text = "x"
    # Mark the top-left cell a vMerge CONTINUATION with no origin row above → invalid grid.
    tc_pr = bad.cell(0, 0)._tc.get_or_add_tcPr()
    tc_pr.append(tc_pr.makeelement(qn("w:vMerge"), {qn("w:val"): "continue"}))
    doc.save(str(p))

    data = ResumeLoader().load(str(p))  # must NOT raise

    assert data.name == "Jane Doe"  # paragraphs still extracted despite the bad table
    assert data.email == "jane@example.com"


def test_load_docx_paragraph_only_resume_unchanged(tmp_path: Path) -> None:
    """Regression: a paragraph-only résumé is unaffected by the table-extraction addition."""
    from docx import Document

    p = tmp_path / "plain.docx"
    doc = Document()
    doc.add_paragraph("John Smith")
    doc.add_paragraph("john@example.com")
    doc.add_paragraph("Experience")
    doc.save(str(p))

    data = ResumeLoader().load(str(p))

    assert data.name == "John Smith"
    assert data.email == "john@example.com"
