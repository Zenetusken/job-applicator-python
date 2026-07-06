#!/usr/bin/env python3
"""Deterministic critical-paths QA / sanity harness for the job-applicator CLI.

Runs the SAFE critical paths through the real CLI as a user would, asserts STABLE
signals (exit codes, no-traceback, valid JSON, artifacts, voice-tells), and emits a
markdown PASS/FAIL/WARN report. Use at the end of an implementation arc and for
general sanity checks. See SKILL.md.

THREE invariants make this trustworthy:

1. ISOLATION. Every CLI call runs with HOME pointed at a throwaway temp dir, so the
   tool's real state dir (~/.job-applicator: dedupe DB, batch progress, cookies,
   authenticated browser-profile) is NEVER touched. Read-only caches (the embedding
   model, the Playwright browser) are shared via HF_HOME / PLAYWRIGHT_BROWSERS_PATH so
   the run stays fast. A guard aborts if isolation isn't in effect, and the real DB is
   fingerprinted before/after as proof. (This harness exists because a non-isolated QA
   run once deleted that DB.)

2. ASSERT CORRECT BEHAVIOR, never the bug. Known bugs are listed in KNOWN_FAIL and
   each asserts what SHOULD happen. A known bug shows as XFAIL (expected, not a
   regression). When a fix lands, the check flips to XPASS — the harness's own signal
   to move that name out of KNOWN_FAIL. A check expected to pass that fails is a FAIL
   (real regression) and exits non-zero.

3. ACCOUNT SAFETY. NEVER runs login / import-cookies / search / apply / check-session
   (they touch the real account or launch a real browser+session). Those are probed
   for --help only. Cover letters use inline --description (no scraping/browser).

LLM output is non-deterministic, so LLM checks assert stable signals (exit 0,
no-traceback, artifact written, voice-tells count as a WARN metric) — never exact text.
The isolated run uses CONFIG DEFAULTS (no real config.toml): reproducible, but it will
not catch config-specific issues.
"""

from __future__ import annotations

import json
import os
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import time
from pathlib import Path

REPO = Path(__file__).resolve().parents[3]
JA = REPO / ".venv" / "bin" / "job-applicator"
VLLM_URL = "http://localhost:8000/v1/models"

# --- Known bugs: each check below asserts CORRECT behavior; these names are expected
# --- to FAIL today (-> XFAIL). When one flips to XPASS, the bug is fixed: delete it here.
# --- EMPTY as of 2026-06-22: the entire QA-arc backlog is cleared. A new entry here marks a
# --- freshly-triaged bug whose check asserts the FIX (XFAIL until fixed → XPASS → promote).
KNOWN_FAIL: set[str] = set()
# Cleared this arc (all promoted XPASS→PASS): ats-check tracebacks; runtime errors→stderr;
# tailor --yes non-interactive; --json --verbose valid JSON; generate-cover-letter litellm-noise;
# config-init bad-path clean error; import-cookies [browser] markup; match skill-match threshold.

# ---------------------------------------------------------------- isolation
REAL_HOME = Path(os.environ["HOME"])
REAL_DB = REAL_HOME / ".job-applicator" / "applications.db"
HF_HOME = os.environ.get("HF_HOME") or str(REAL_HOME / ".cache" / "huggingface")
PW_PATH = os.environ.get("PLAYWRIGHT_BROWSERS_PATH") or str(REAL_HOME / ".cache" / "ms-playwright")

WORK = Path(tempfile.mkdtemp(prefix="ja-qa-"))
ISO_HOME = WORK / "home"
ISO_HOME.mkdir(parents=True)
ENV = {
    **os.environ,
    "HOME": str(ISO_HOME),
    "HF_HOME": HF_HOME,
    "PLAYWRIGHT_BROWSERS_PATH": PW_PATH,
    "NO_COLOR": "1",
    "COLUMNS": "200",
}
# NO_COLOR alone is NOT enough under a dev shell's FORCE_COLOR: FORCE_COLOR forces Rich's
# is_terminal=True, so the progress SPINNER still ANIMATES onto a piped stdout (NO_COLOR only strips
# color codes, not the animation). A user piping `--json` has neither var set, so scrub the
# color-FORCERS to make the harness faithful to a real pipe (else --json checks see spurious ANSI).
for _cvar in ("FORCE_COLOR", "CLICOLOR_FORCE", "PY_COLORS"):
    ENV.pop(_cvar, None)

# GUARD: refuse to run if isolation isn't real — a broken edit must not touch real state.
if Path(ENV["HOME"]).resolve() == REAL_HOME.resolve() or str(WORK) not in ENV["HOME"]:
    sys.exit("FATAL: isolation not in effect — refusing to run (would touch real ~/.job-applicator)")
if not JA.exists():
    sys.exit(f"FATAL: {JA} not found — run setup (see SKILL.md)")


def _db_sig(p: Path) -> tuple[float, int] | None:
    return (p.stat().st_mtime, p.stat().st_size) if p.exists() else None


REAL_DB_BEFORE = _db_sig(REAL_DB)

# ---------------------------------------------------------------- results
results: list[dict[str, str]] = []


def record(name: str, tier: str, ok: bool, detail: str = "") -> bool:
    if name in KNOWN_FAIL:
        status = "XPASS" if ok else "XFAIL"  # XPASS = known bug now fixed → promote it
    else:
        status = "PASS" if ok else "FAIL"
    results.append({"name": name, "tier": tier, "status": status, "detail": detail})
    return ok


def warn(name: str, tier: str, detail: str) -> None:
    results.append({"name": name, "tier": tier, "status": "WARN", "detail": detail})


def skip(name: str, tier: str, detail: str) -> None:
    results.append({"name": name, "tier": tier, "status": "SKIP", "detail": detail})


def run(
    *args: str,
    timeout: int = 120,
    stdin: str | None = None,
    extra_env: dict[str, str] | None = None,
) -> subprocess.CompletedProcess[str]:
    """Run the CLI isolated, from the temp workdir. Never raises on non-zero exit.

    ``extra_env`` overlays ENV for one call — used to deliberately FORCE color (the opposite of the
    harness default) to prove ``--json`` stays pure regardless of the caller's color env.
    """
    try:
        return subprocess.run(
            [str(JA), *args],
            env={**ENV, **(extra_env or {})},
            cwd=str(WORK),
            capture_output=True,
            text=True,
            timeout=timeout,
            input=stdin,
        )
    except subprocess.TimeoutExpired as exc:
        out = (exc.stdout or b"").decode() if isinstance(exc.stdout, bytes) else (exc.stdout or "")
        err = (exc.stderr or b"").decode() if isinstance(exc.stderr, bytes) else (exc.stderr or "")
        return subprocess.CompletedProcess(args, 124, out, err + "\n<TIMEOUT>")


def has_traceback(cp: subprocess.CompletedProcess[str]) -> bool:
    return "Traceback (most recent call last)" in (cp.stdout + cp.stderr)


def grounding_fail_closed(cp: subprocess.CompletedProcess[str]) -> bool:
    combined = (cp.stdout + cp.stderr).lower()
    return (
        cp.returncode != 0
        and not has_traceback(cp)
        and "grounding verification" in combined
    )


def cover_letter_fail_closed(cp: subprocess.CompletedProcess[str]) -> bool:
    """A requested cover letter can fail closed at either honesty layer."""
    combined = (cp.stdout + cp.stderr).lower()
    deterministic_guard = (
        "generated cover letter" in combined
        and (
            "unearned" in combined
            or "not in the résumé" in combined
            or "falsely claims" in combined
            or "placeholder text" in combined
            or "proper sign-off" in combined
        )
    )
    return grounding_fail_closed(cp) or (
        cp.returncode != 0 and not has_traceback(cp) and deterministic_guard
    )


def vllm_up() -> bool:
    try:
        import urllib.request

        with urllib.request.urlopen(VLLM_URL, timeout=3) as r:  # noqa: S310
            return r.status == 200
    except Exception:
        return False


# ---------------------------------------------------------------- fixtures
def make_fixtures() -> dict[str, Path]:
    from docx import Document  # core dep

    rdoc = WORK / "resume.docx"
    d = Document()
    d.add_heading("Jordan Sample", 0)
    d.add_paragraph("jordan.sample@example.com | (555) 123-4567 | San Francisco, CA")
    for h, lines in [
        ("Summary", ["Senior Python engineer, 8 years building async data pipelines and ML services."]),
        ("Experience", ["Staff Engineer, Acme Data (2021-Present)",
                        "Built async ingestion handling 2B events/day with asyncio.",
                        "Led a Pydantic v2 + mypy-strict migration across 40 services.",
                        "Backend Engineer, Globex (2017-2021)",
                        "Designed PostgreSQL schemas and Redis caching for a 5M-user app."]),
        ("Education", ["B.S. Computer Science, State University (2017)"]),
        # Deliberately a comma list WRAPPED across two lines — the realistic résumé shape
        # that exposed the F-A skill-parse bug (a single line parsed fine and hid it). Keep
        # it multi-line so the harness exercises the real case, not the easy one.
        ("Skills", ["Python, asyncio, FastAPI, Pydantic,",
                    "PostgreSQL, Redis, Docker, AWS"]),
    ]:
        d.add_heading(h, 1)
        for ln in lines:
            d.add_paragraph(ln)
    d.save(str(rdoc))

    low = WORK / "lowqual.docx"
    dl = Document()
    dl.add_paragraph("Jordan. I do python things. Hire me.")
    dl.save(str(low))

    corrupt = WORK / "corrupt.pdf"
    corrupt.write_bytes(b"not a real pdf %%\x00\x01garbage")
    empty = WORK / "empty.docx"
    empty.write_bytes(b"")

    jobs = [
        {"title": "Senior Python Engineer", "company": "Initech",
         "url": "https://example.com/jobs/1",
         "description": "Async data pipelines in Python; asyncio, Pydantic, PostgreSQL, AWS.",
         "requirements": ["Python", "asyncio", "PostgreSQL"], "board": "linkedin"},
        {"title": "Backend Engineer", "company": "Globex",
         "url": "https://example.com/jobs/2",
         "description": "High-throughput async services; asyncio, Redis, PostgreSQL.",
         "requirements": ["Python", "Redis"], "board": "linkedin"},
        {"title": "Frontend React Developer", "company": "WebStart",
         "url": "https://example.com/jobs/3",
         "description": "Build React UIs with TypeScript.",
         "requirements": ["React", "TypeScript"], "board": "linkedin"},
    ]
    jf = WORK / "jobs.json"
    jf.write_text(json.dumps(jobs))
    bigjobs = jobs + [
        {"title": f"Engineer {i}", "company": f"Co{i}", "url": f"https://example.com/jobs/1{i}",
         "description": "Python, asyncio, AWS, distributed systems.",
         "requirements": ["Python", "AWS"], "board": "linkedin"}
        for i in range(5)
    ]
    bf = WORK / "bigjobs.json"
    bf.write_text(json.dumps(bigjobs))
    malformed = WORK / "malformed.json"
    malformed.write_text('[{"title":"X",')
    return {"resume": rdoc, "low": low, "corrupt": corrupt, "empty": empty,
            "jobs": jf, "bigjobs": bf, "malformed": malformed}


# ---------------------------------------------------------------- CORE checks (offline)
def core_checks(fx: dict[str, Path]) -> None:
    t = "CORE"

    cp = run("--version")
    record("global: --version", t, cp.returncode == 0 and "job-applicator v" in cp.stdout,
            f"exit={cp.returncode}")

    cp = run("--help")
    record("global: --help lists commands", t, cp.returncode == 0 and "doctor" in cp.stdout,
            f"exit={cp.returncode}")

    cp = run("frobnicate")
    record("global: unknown command → exit 2", t, cp.returncode == 2, f"exit={cp.returncode}")

    cp = run("--log-file", "x.log", "ats-check", "--resume", str(fx["resume"]))
    record("global: --log-file without --verbose rejected", t, cp.returncode == 2,
           f"exit={cp.returncode}")

    cp = run("config-init", "-o", str(WORK / "c.toml"))
    record("config-init: writes sample", t, cp.returncode == 0 and (WORK / "c.toml").exists(),
           f"exit={cp.returncode}")

    # Regression guard (#43): a bad output path is a clean error, not a traceback.
    cp = run("config-init", "-o", "/proc/nope/c.toml")
    record("config-init: bad output path → clean error (no traceback)", t,
           cp.returncode != 0 and not has_traceback(cp), f"exit={cp.returncode}")

    # B6 (QA pass #2): a bad --jobs-file must be a clean typed error, never a raw traceback.
    # match had NO inner handler → its outer `except Exception` re-raised → Typer traceback.
    # Fails at the jobs-load (before embeddings/tailoring) → offline/CORE.
    notalist = WORK / "jobs_notalist.json"
    notalist.write_text("{}")
    badfields = WORK / "jobs_badfields.json"
    badfields.write_text('[{"title": "x"}]')
    jobsdir = WORK / "jobsdir"  # a directory path → IsADirectoryError (OSError, not FileNotFound)
    jobsdir.mkdir(exist_ok=True)
    nonutf8 = WORK / "jobs_nonutf8.json"  # non-UTF-8 bytes → UnicodeDecodeError (a ValueError)
    nonutf8.write_bytes(b"\xff\xfe not utf-8 \x80\x81")
    for label, jf2 in (("malformed", fx["malformed"]), ("not-a-list", notalist),
                       ("missing-fields", badfields), ("missing-file", WORK / "nope.json"),
                       ("directory", jobsdir), ("non-utf8", nonutf8)):
        for cmd in ("match", "batch"):
            extra = ("--no-cover-letter",) if cmd == "batch" else ()
            cp = run(cmd, "--resume", str(fx["resume"]), "--jobs-file", str(jf2), *extra)
            record(f"{cmd}: bad jobs-file ({label}) → clean error, no traceback", t,
                   cp.returncode != 0 and not has_traceback(cp), f"exit={cp.returncode}")

    cp = run("ats-check", "--resume", str(fx["resume"]))
    record("ats-check: valid docx", t, cp.returncode == 0 and "ATS Score" in cp.stdout,
           f"exit={cp.returncode}")

    cp = run("ats-check", "--resume", str(fx["resume"]), "--json")
    ok_json = False
    try:
        ok_json = isinstance(json.loads(cp.stdout)["score"], (int, float))
    except Exception:
        pass
    record("ats-check: --json is valid JSON", t, cp.returncode == 0 and ok_json,
           f"exit={cp.returncode}")

    # Bad-input cases must produce a CLEAN error: nonzero exit, no traceback, and NOT a
    # timeout (124). --ocr-mode off on the corrupt pdf avoids the slow PaddleOCR fallback
    # (whose model cache isn't shared into the isolated HOME) so the check stays fast.
    for label, path, extra in [("missing file", str(WORK / "nope.docx"), []),
                               ("corrupt pdf", str(fx["corrupt"]), ["--ocr-mode", "off"]),
                               ("empty file", str(fx["empty"]), []),
                               ("directory path", str(WORK), [])]:
        cp = run("ats-check", "--resume", path, *extra, timeout=40)
        record(f"ats-check: {label} → clean error (no traceback)", t,
               cp.returncode not in (0, 124) and not has_traceback(cp), f"exit={cp.returncode}")

    # Runtime errors must go to stderr (the tool's own stdout=data / stderr=logs contract).
    cp = run("ats-check")  # no --resume
    record("ats-check: no-resume error goes to stderr (not stdout)", t,
           "Resume path required" in cp.stderr and "Resume path required" not in cp.stdout,
           f"on_stdout={'Resume path required' in cp.stdout}")

    # QA pass #2 — input validation / honest failures (all offline/CORE):
    blanktxt = WORK / "blank.txt"
    blanktxt.write_text("   \n")  # a VALID file with no extractable text (cf. empty.docx = bad bytes)
    cp = run("ats-check", "--resume", str(blanktxt))
    record("B3 ats-check: empty/no-text résumé → clean 'no extractable text' error", t,
           cp.returncode != 0 and not has_traceback(cp) and "no extractable text" in cp.stderr,
           f"exit={cp.returncode}")
    cp = run("ats-check", "--resume", str(fx["resume"]), "--ocr-mode", "bogus")
    record("B2 --ocr-mode invalid value → rejected (exit 2)", t, cp.returncode == 2,
           f"exit={cp.returncode}")
    cp = run("match", "--resume", str(fx["resume"]))  # no --jobs-file
    record("B7 match: no jobs source → clean error (not silent demo jobs)", t,
           cp.returncode != 0 and not has_traceback(cp) and "Provide --jobs-file" in cp.stderr,
           f"exit={cp.returncode}")
    cp = run("match", "--resume", str(fx["resume"]), "--jobs-file", str(fx["jobs"]), "--top-k", "0")
    record("B8 match: --top-k 0 (out of range) → rejected (exit 2)", t, cp.returncode == 2,
           f"exit={cp.returncode}")
    cp = run("match", "--resume", str(fx["resume"]), "--jobs-file", str(fx["jobs"]),
             "--min-score", "2.0")
    record("B8 match: --min-score >1 → rejected (exit 2)", t, cp.returncode == 2,
           f"exit={cp.returncode}")
    cp = run("config-init", "-o", str(WORK))  # a directory, not a file
    record("B4 config-init: -o <directory> → clean error (not 'already exists')", t,
           cp.returncode != 0 and not has_traceback(cp) and "directory" in cp.stderr,
           f"exit={cp.returncode}")

    # import-cookies help renders "[browser]" as Rich markup and EATS it ("the  extra").
    cp = run("import-cookies", "--help")
    record("import-cookies: --help shows the [browser] extra name (markup not eaten)", t,
           "the  extra" not in cp.stdout, "double-space gap present" if "the  extra" in cp.stdout else "ok")

    # ats-check --strict gates the exit on the verdict; without it, an incompatible résumé is
    # report-only (exit 0). (PR A — was a WARN.)
    cp_def = run("ats-check", "--resume", str(fx["low"]))
    cp_strict = run("ats-check", "--resume", str(fx["low"]), "--strict")
    record("ats-check: --strict exits non-zero on an incompatible résumé", t,
           cp_def.returncode == 0 and cp_strict.returncode != 0,
           f"default={cp_def.returncode} strict={cp_strict.returncode}")

    # --log-file to an unwritable path WARNS (verbose.py guard) and is non-fatal by design — a
    # failed diagnostic log shouldn't fail an otherwise-successful command. (PR A — was a WARN.)
    cp = run("--verbose", "--log-file", "/proc/nope/x.log", "ats-check", "--resume", str(fx["resume"]))
    record("--log-file unwritable → warns (non-fatal, by design)", t,
           cp.returncode == 0 and "Could not write verbose log" in (cp.stdout + cp.stderr),
           f"exit={cp.returncode}")

    # Regression guard (#42): the --verbose report goes to stderr, so --json stdout stays JSON.
    cp = run("ats-check", "--resume", str(fx["resume"]), "--json", "--verbose")
    ok_jv = False
    try:
        json.loads(cp.stdout)
        ok_jv = True
    except Exception:
        pass
    record("global: --json --verbose → still valid JSON", t, ok_jv, f"exit={cp.returncode}")

    # Runtime errors → stderr is fleet-wide (shared err_console); gate a sibling too, not
    # just ats-check. `batch` with no jobs source errors before any vLLM/scrape (offline).
    cp = run("batch", "--resume", str(fx["resume"]), "--no-cover-letter")
    record("batch: error message goes to stderr (not stdout)", t,
           cp.returncode != 0
           and "Provide --jobs-file" in cp.stderr
           and "Provide --jobs-file" not in cp.stdout,
           f"exit={cp.returncode}")

    for cmd in ("apply", "login", "import-cookies"):
        cp = run(cmd, "--help")
        record(f"account-safe: {cmd} --help works (NOT executed)", t,
               cp.returncode == 0 and "Usage" in cp.stdout, f"exit={cp.returncode}")

    # --- Funnel backbone (Increment 1): the `status` dashboard + job store, offline.
    # apply --from / saved-list stay account-safe (probed --help-only above); their
    # gating is unit-covered (test_funnel_cli asserts NO browser is built on a bad/empty
    # target). Here we exercise the offline read path end-to-end through the real CLI.
    cp = run("status")
    record("status: runs on an empty store (clean, no traceback)", t,
           cp.returncode == 0 and "Funnel" in cp.stdout and not has_traceback(cp),
           f"exit={cp.returncode}")
    cp = run("status", "--json")
    ok_sj = False
    try:
        payload = json.loads(cp.stdout)
        ok_sj = payload["total"] == 0 and {"found", "matched", "applied"} <= set(payload["counts"])
    except Exception:
        pass
    record("status: --json is valid JSON with funnel counts", t, cp.returncode == 0 and ok_sj,
           f"exit={cp.returncode}")

    # Seed one job straight into the isolated store (status above created the schema), then
    # prove `status` reflects it — the store→status read path, offline & account-safe.
    seeded_ok = False
    try:
        c = sqlite3.connect(str(ISO_HOME / ".job-applicator" / "applications.db"))
        c.execute(
            "INSERT INTO jobs (job_url, title, company, board, funnel_status, "
            "first_seen_at, updated_at) VALUES (?,?,?,?,?,?,?)",
            ("https://example.com/qa/seed", "QA Seeded Eng", "QACo", "linkedin", "found",
             "2026-06-22T00:00:00+00:00", "2026-06-22T00:00:00+00:00"),
        )
        c.commit()
        c.close()
        cp = run("status", "--json")
        payload = json.loads(cp.stdout)
        seeded_ok = payload["total"] == 1 and payload["counts"]["found"] == 1
    except Exception:
        pass
    record("status: reflects a job in the store (read path)", t, seeded_ok, "store→status")

    # TUI launch wiring (Increment 2). The harness captures stdout (NOT a TTY), so the
    # full-screen app never actually launches here — bare invocation prints help, `tui`
    # errors cleanly. The interactive app itself is covered by Pilot unit tests.
    cp = run()  # bare invocation, non-TTY
    record("tui: bare invocation (non-TTY) → help, exit 0, no traceback", t,
           cp.returncode == 0 and "Usage" in cp.stdout and not has_traceback(cp),
           f"exit={cp.returncode}")
    cp = run("tui", "--help")
    record("tui: --help works", t, cp.returncode == 0 and "Usage" in cp.stdout,
           f"exit={cp.returncode}")
    cp = run("tui")  # non-TTY → clean guard, never a Textual launch
    record("tui: non-TTY → clean error (not a crash)", t,
           cp.returncode == 1 and not has_traceback(cp), f"exit={cp.returncode}")


# ---------------------------------------------------------------- LIVE checks (need vLLM)
def live_checks(fx: dict[str, Path]) -> None:
    t = "LIVE"

    cp = run("doctor", timeout=180)
    record("doctor: runs, LLM reachable", t,
           cp.returncode == 0 and "reachable" in cp.stdout, f"exit={cp.returncode}")
    cp = run("doctor", "--json", timeout=180)
    ok_dj = False
    try:
        ok_dj = "ok" in json.loads(cp.stdout)
    except Exception:
        pass
    record("doctor: --json emits valid JSON (incl. ok verdict)", t, ok_dj, f"exit={cp.returncode}")

    cp = run("match", "--resume", str(fx["resume"]), "--jobs-file", str(fx["jobs"]), "--top-k", "3",
             timeout=180)
    scraped = any(s in (cp.stdout + cp.stderr).lower() for s in ("playwright", "chromium", "cloudflare"))
    record("match: --jobs-file ranks (no scraping)", t,
           cp.returncode == 0 and "Score" in cp.stdout and not scraped, f"exit={cp.returncode}")

    cp = run("match", "--resume", str(fx["resume"]), "--jobs-file", str(fx["jobs"]), "--json",
             timeout=180)
    arr = None
    try:
        arr = json.loads(cp.stdout)
    except Exception:
        pass
    ok_mj = isinstance(arr, list) and bool(arr) and "title" in arr[0] and "score" in arr[0]
    record("match: --json is valid JSON", t, cp.returncode == 0 and ok_mj, f"exit={cp.returncode}")

    # Adversarial stdout-contract guard: --json must stay PURE JSON even when the caller FORCES
    # color (CI, a TTY, FORCE_COLOR). Proves the progress spinner is on stderr, not stdout — FAILS
    # on the pre-fix code (console.status spinner → stdout) and passes once it's on err_console.
    cp_fc = run("match", "--resume", str(fx["resume"]), "--jobs-file", str(fx["jobs"]), "--json",
                timeout=180, extra_env={"FORCE_COLOR": "3"})
    fc_ok = False
    try:
        fc_ok = isinstance(json.loads(cp_fc.stdout), list)
    except Exception:
        pass
    record("match: --json stays pure JSON under FORCE_COLOR (spinner on stderr)", t, fc_ok,
           "stdout not pure JSON — a progress spinner is leaking to stdout" if not fc_ok else "ok")

    # Regression guard (#45): at the 0.75 skill-match threshold a Python-only résumé reports
    # React/TypeScript as MISSING for a React job (the old 0.55 wrongly "covered" them).
    react_ok = False
    if isinstance(arr, list):
        react = next((j for j in arr if "React" in j.get("title", "")), None)
        if react is not None:
            miss = " ".join(react.get("missing_skills", [])).lower()
            react_ok = "react" in miss or "typescript" in miss
    record("match: React job reports React/TypeScript as missing skills", t, react_ok,
           "missing_skills under-reports the gap" if not react_ok else "ok")

    # QA F-C (blind-spot guard): the React→missing check above is DEGENERATE — it passes even
    # when matching is fully broken, because an empty matched set makes every requirement
    # trivially "missing". Assert the POSITIVE direction too: a strong-overlap job must report a
    # skill the résumé actually HAS as MATCHED. Combined with the wrapped multi-line Skills
    # fixture, this would have caught both the F-A parser blob bug and the F-B grounding drop.
    matched_ok = False
    if isinstance(arr, list):
        strong = next((j for j in arr if "Python Engineer" in j.get("title", "")), None)
        if strong is not None:
            got = " ".join(strong.get("matched_skills", [])).lower()
            matched_ok = any(s in got for s in ("python", "asyncio", "postgresql"))
    record("match: strong-overlap job reports a résumé skill as MATCHED (not just missing)", t,
           matched_ok,
           "matched_skills empty — résumé's own skills not recognized" if not matched_ok else "ok")

    # Increment 1: match persists scored jobs into the funnel store (a side effect of the
    # LIVE match runs above) — assert at the isolated DB directly.
    fdb = ISO_HOME / ".job-applicator" / "applications.db"
    jobs_rows = 0
    if fdb.exists():
        fc = sqlite3.connect(str(fdb))
        try:
            jobs_rows = fc.execute("select count(*) from jobs").fetchone()[0]
        except sqlite3.Error:
            jobs_rows = 0
        fc.close()
    record("match: persists scored jobs to the store", t, jobs_rows >= 1, f"jobs_rows={jobs_rows}")

    cp = run("generate-cover-letter", "-t", "Senior Python Engineer", "-c", "Initech",
             "-d", "Async pipelines; asyncio, Pydantic, PostgreSQL, AWS.",
             "--resume", str(fx["resume"]), timeout=240)
    letter = cp.stdout.split("Generated Cover Letter:", 1)[-1]
    record("generate-cover-letter: inline produces or fails closed on validation", t,
           (cp.returncode == 0 and "Generated Cover Letter" in cp.stdout)
           or cover_letter_fail_closed(cp), f"exit={cp.returncode}")
    # KNOWN: litellm framework noise leaks to stdout on SUCCESS (instructor tool-call path
    # always fails on this vLLM → fallback → banner), polluting a redirected letter.
    record("generate-cover-letter: output free of litellm framework noise", t,
           "Give Feedback" not in cp.stdout and "BadRequestError" not in cp.stdout,
           "litellm banner on stdout" if "Give Feedback" in cp.stdout else "clean")
    # voice-tells: a METRIC, not a hard gate (LLM output flaps); WARN if elevated.
    try:
        sys.path.insert(0, str(REPO / "src"))
        from job_applicator.documents.cover_letter import CoverLetterGenerator

        tells = CoverLetterGenerator._voice_tells(letter)
        md = "`" in letter
        if md or len(tells) > 2:
            warn("generate-cover-letter: voice quality", t,
                 f"voice_tells={tells}; markdown_leak={md}")
        else:
            record("generate-cover-letter: voice clean (markdown-free, ≤2 tells)", t, True,
                   f"voice_tells={tells}")
    except Exception as exc:  # pragma: no cover
        warn("generate-cover-letter: voice quality", t, f"could not score: {exc}")

    cp = run("generate-cover-letter", "-t", "Senior Python Engineer", "-c", "Initech",
             "-d", "Async pipelines; asyncio, Pydantic.", "--resume", str(fx["resume"]),
             "--json", timeout=240)
    ok_gj = False
    try:
        ok_gj = bool(json.loads(cp.stdout).get("cover_letter"))
    except Exception:
        pass
    record("generate-cover-letter: --json emits JSON or fails closed on validation", t,
           ok_gj or cover_letter_fail_closed(cp),
           f"exit={cp.returncode}")

    # Regression guard (#39): --yes must be non-interactive (it used to hang on the action menu).
    cp = run("tailor", "-t", "Senior Python Engineer", "-c", "Initech",
             "-d", "Async pipelines; asyncio, Pydantic, PostgreSQL.",
             "--resume", str(fx["resume"]), "--yes", timeout=200)
    out_dir = WORK / "output"
    wrote = out_dir.exists() and any(out_dir.glob("tailored_*.txt"))
    record("tailor: --yes is non-interactive (exits, writes artifact)", t,
           cp.returncode == 0 and wrote, f"exit={cp.returncode} (124=hang)")
    cp = run("tailor", "-t", "Senior Python Engineer", "-c", "Initech",
             "-d", "Async pipelines; asyncio, Pydantic.", "--resume", str(fx["resume"]),
             "--json", timeout=200)
    ok_tj = False
    try:
        ok_tj = bool(json.loads(cp.stdout).get("tailored_text"))
    except Exception:
        pass
    record("tailor: --json emits valid JSON (clean stdout, implies --yes)", t, ok_tj,
           f"exit={cp.returncode}")

    cp = run("tailor", "-t", "Chef", "-c", "Restaurant", "-d", "Cook food in a kitchen.",
             "--resume", str(fx["resume"]), "--yes", "--min-score", "0.99", timeout=180)
    record("tailor: --min-score abort path", t,
           cp.returncode == 0 and "Aborting" in cp.stdout, f"exit={cp.returncode}")

    # Increment 1: `tailor --from <url>` tailors a stored job (no -t/-c retyping) and marks
    # it tailored in the store — the funnel head→tailor handoff, end-to-end.
    tfrom_url = "https://example.com/qa/tfrom"
    tfrom_ok = False
    tfrom_exit = "n/a"
    try:
        tc = sqlite3.connect(str(ISO_HOME / ".job-applicator" / "applications.db"))
        tc.execute(
            "INSERT OR IGNORE INTO jobs (job_url, title, company, board, funnel_status, "
            "description, first_seen_at, updated_at) VALUES (?,?,?,?,?,?,?,?)",
            (tfrom_url, "Senior Python Engineer", "Initech", "linkedin", "found",
             "Async pipelines; asyncio, Pydantic, PostgreSQL.",
             "2026-06-22T00:00:00+00:00", "2026-06-22T00:00:00+00:00"),
        )
        tc.commit()
        tc.close()
        cp = run("tailor", "--from", tfrom_url, "--resume", str(fx["resume"]), "--yes", timeout=200)
        tfrom_exit = str(cp.returncode)
        tc = sqlite3.connect(str(ISO_HOME / ".job-applicator" / "applications.db"))
        row = tc.execute("select funnel_status from jobs where job_url=?", (tfrom_url,)).fetchone()
        tc.close()
        tfrom_ok = cp.returncode == 0 and row is not None and row[0] in ("tailored", "cover_letter")
    except Exception as exc:
        tfrom_exit = f"err:{exc}"
    record("tailor: --from a stored job tailors it (marks it tailored)", t, tfrom_ok,
           f"exit={tfrom_exit}")

    # batch: clean multi-job run + artifacts
    cp = run("batch", "--resume", str(fx["resume"]), "--jobs-file", str(fx["jobs"]),
             "--no-cover-letter", "--run-id", "qa_a", timeout=300)
    bscraped = any(s in (cp.stdout + cp.stderr).lower() for s in ("playwright", "chromium"))
    record("batch: --jobs-file tailors (no scraping)", t,
           cp.returncode == 0 and "tailored" in cp.stdout and not bscraped, f"exit={cp.returncode}")

    cp = run("batch", "--resume", str(fx["resume"]), "--jobs-file", str(fx["malformed"]),
             "--no-cover-letter", timeout=60)
    record("batch: malformed jobs → clean error (no traceback)", t,
           cp.returncode != 0 and not has_traceback(cp), f"exit={cp.returncode}")

    crash_recovery_check(fx, t)


def crash_recovery_check(fx: dict[str, Path], t: str) -> None:
    """Deterministic resume test: run a batch, then SIMULATE an interrupted run by
    editing the isolated DB (mark run incomplete, drop 3 completed jobs + artifacts),
    and assert --resume-run skips the survivors and finishes the rest. No kill-timing
    flakiness — the resume LOGIC is what matters and the units cover the kill path."""
    db = ISO_HOME / ".job-applicator" / "applications.db"
    cp = run("batch", "--resume", str(fx["resume"]), "--jobs-file", str(fx["bigjobs"]),
             "--top-k", "8", "--no-cover-letter", "--run-id", "qa_cr", timeout=420)
    if cp.returncode != 0 or not db.exists():
        skip("batch: crash recovery (resume skips completed)", t,
             f"setup run failed (exit={cp.returncode}); resume logic covered by test_batch_state")
        return
    try:
        c = sqlite3.connect(str(db))
        done = c.execute(
            "select count(*) from batch_jobs where status='completed' and run_id='qa_cr'"
        ).fetchone()[0]
        if done < 5:
            c.close()
            skip("batch: crash recovery (resume skips completed)", t,
                 f"only {done} completed; need ≥5 to simulate partial")
            return
        victims = [r[0] for r in c.execute(
            "select job_url from batch_jobs where status='completed' and run_id='qa_cr' limit 3")]
        c.execute(
            "delete from batch_jobs where run_id='qa_cr' and job_url in (?,?,?)", victims)
        c.execute("update batch_runs set status='running' where run_id='qa_cr'")
        c.commit()
        c.close()
    except Exception as exc:
        skip("batch: crash recovery (resume skips completed)", t, f"db edit failed: {exc}")
        return
    cp = run("batch", "--resume", str(fx["resume"]), "--jobs-file", str(fx["bigjobs"]),
             "--top-k", "8", "--no-cover-letter", "--run-id", "qa_cr", "--resume-run", timeout=420)
    c = sqlite3.connect(str(db))
    final = c.execute(
        "select count(*) from batch_jobs where status='completed' and run_id='qa_cr'"
    ).fetchone()[0]
    c.close()
    record("batch: crash recovery (resume skips completed)", t,
           cp.returncode == 0 and "Resuming" in cp.stdout and "Skipping" in cp.stdout and final == 8,
           f"exit={cp.returncode}, final_completed={final}")


# ---------------------------------------------------------------- report
def render() -> int:
    counts: dict[str, int] = {}
    for r in results:
        counts[r["status"]] = counts.get(r["status"], 0) + 1
    real_after = _db_sig(REAL_DB)
    isolation_ok = REAL_DB_BEFORE == real_after

    lines = ["# job-applicator — QA sanity report", ""]
    lines.append(f"isolation: real ~/.job-applicator/applications.db **{'UNCHANGED ✓' if isolation_ok else 'CHANGED ✗ — ISOLATION BREACH'}**")
    lines.append("note: isolated run uses CONFIG DEFAULTS (no real config.toml) — reproducible, "
                 "won't catch config-specific issues.")
    lines.append("")
    summary = "  ".join(f"{k}={counts[k]}" for k in
                        ("PASS", "FAIL", "XFAIL", "XPASS", "WARN", "SKIP") if k in counts)
    lines.append(f"**{summary}**")
    lines.append("")
    lines.append("| Status | Tier | Check | Detail |")
    lines.append("|---|---|---|---|")
    order = {"FAIL": 0, "XPASS": 1, "WARN": 2, "XFAIL": 3, "SKIP": 4, "PASS": 5}
    icon = {"PASS": "✓ PASS", "FAIL": "✗ FAIL", "XFAIL": "• XFAIL", "XPASS": "▲ XPASS",
            "WARN": "! WARN", "SKIP": "– SKIP"}
    for r in sorted(results, key=lambda r: (order.get(r["status"], 9), r["tier"])):
        lines.append(f"| {icon[r['status']]} | {r['tier']} | {r['name']} | {r['detail']} |")
    lines.append("")
    if counts.get("XFAIL"):
        lines.append("_XFAIL = known open bug, expected to fail (not a regression)._")
    if counts.get("XPASS"):
        lines.append("_XPASS = a known bug now PASSES — fixed! Remove its name from KNOWN_FAIL in qa.py._")
    if not isolation_ok:
        lines.append("_ISOLATION BREACH: the run mutated real state. Stop and fix qa.py before reuse._")

    report = "\n".join(lines)
    out = Path("/tmp/job-applicator-qa-report.md")
    out.write_text(report + "\n")
    print(report)
    print(f"\n(report written to {out})")

    regressions = counts.get("FAIL", 0)
    return 1 if (regressions or not isolation_ok) else 0


def main() -> int:
    mode = sys.argv[1] if len(sys.argv) > 1 else "all"
    print(f"qa-sanity: repo={REPO}  isolated HOME={ISO_HOME}\n", file=sys.stderr)
    try:
        fx = make_fixtures()
        if mode in ("all", "--core"):
            core_checks(fx)
        if mode in ("all", "--live"):
            if vllm_up():
                live_checks(fx)
            else:
                skip("LIVE tier", "LIVE", "vLLM unreachable at localhost:8000")
        return render()
    finally:
        shutil.rmtree(WORK, ignore_errors=True)


if __name__ == "__main__":
    raise SystemExit(main())
